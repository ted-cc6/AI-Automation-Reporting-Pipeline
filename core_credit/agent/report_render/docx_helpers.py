"""Low-level, reusable formatting primitives -- each one applies brand.py's constants directly,
so section_layout.py's per-Part rendering code never touches a hex colour or a font name itself.
Mirrors the skill's own "Formatting Rules" and "Checklist Before Finalising" section by section.
"""

from __future__ import annotations

from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from . import brand, skill_assets


def setup_page(doc: Document) -> None:
    """US Letter, 1-inch margins -- skill's "Page Setup" section."""
    section = doc.sections[0]
    section.page_width = brand.PAGE_WIDTH
    section.page_height = brand.PAGE_HEIGHT
    section.left_margin = brand.PAGE_MARGIN
    section.right_margin = brand.PAGE_MARGIN
    section.top_margin = brand.PAGE_MARGIN
    section.bottom_margin = brand.PAGE_MARGIN


def setup_default_style(doc: Document) -> None:
    """Body text default: Inter (Calibri fallback), Grey 700, 11pt, 1.15 line spacing."""
    style = doc.styles["Normal"]
    style.font.name = brand.FONT_BODY
    style.font.size = brand.SIZE_BODY
    style.font.color.rgb = brand.GREY_700
    style.paragraph_format.line_spacing = brand.LINE_SPACING_BODY
    _set_east_asian_font(style.font, brand.FONT_FALLBACK)


def _set_east_asian_font(font, fallback: str) -> None:
    """python-docx's font.name only sets the Latin-script typeface; without also setting the
    fallback on rFonts' other slots, some Word versions silently substitute a default font for
    a run instead of falling back to Calibri as the skill specifies.
    """
    rpr = font.element
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:cs", "w:hAnsi"):
        rfonts.set(qn(attr), fallback)


def setup_header(doc: Document, logo_variant: str = "primary") -> None:
    """Right-aligned WVI logo in the header, ~3:1 width:height (skill's "Header / Footer")."""
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.add_picture(skill_assets.get_logo_bytes(logo_variant), width=Inches(1.5), height=Inches(0.5))


def setup_footer(doc: Document, run_id: Optional[str] = None, model_version: Optional[str] = None) -> None:
    """Centred page number, Inter 9pt Grey 700 (skill's "Header / Footer"). When run_id/
    model_version are given, prefixes them to the page number so any figure in the document
    can be traced back to the exact run and model that produced it -- added after a real
    incident where the same wave's qualitative theme-tagging produced materially different
    findings across two separate runs, with nothing in the document saying which run's numbers
    a reader was looking at (see CoreCreditImpactReport's own docstring for why bit-for-bit
    reproducibility of that step isn't achievable to begin with).
    """
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    prefix = ""
    if run_id and model_version:
        prefix = f"Run {run_id} ({model_version})  |  "
    elif run_id:
        prefix = f"Run {run_id}  |  "
    elif model_version:
        prefix = f"{model_version}  |  "
    run = paragraph.add_run(f"{prefix}Page " if prefix else "")
    run.font.name = brand.FONT_BODY
    run.font.size = brand.SIZE_CAPTION
    run.font.color.rgb = brand.GREY_700
    _add_page_number_field(paragraph)


def _add_page_number_field(paragraph) -> None:
    """Inserts a real Word PAGE field (updates automatically), not a hardcoded page number."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_divider(doc: Document) -> None:
    """A bottom-border paragraph rule in Orange -- never a table -- per the skill's own
    instruction ("Divider: ... never use a table as a divider").
    """
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FF5515")
    p_borders.append(bottom)
    p_pr.append(p_borders)


def add_title(doc: Document, text: str) -> None:
    run = doc.add_paragraph().add_run(text)
    run.font.name = brand.FONT_HEADING
    run.font.size = brand.SIZE_MAIN_TITLE
    run.font.bold = True
    run.font.color.rgb = brand.MIDNIGHT


def add_subtitle(doc: Document, text: str) -> None:
    run = doc.add_paragraph().add_run(text)
    run.font.name = brand.FONT_HEADING
    run.font.size = brand.SIZE_SUBTITLE
    run.font.color.rgb = brand.GREY_700


def add_section_heading(doc: Document, text: str) -> None:
    """Part-level heading (e.g. "Part 3 -- Business & Household Impact"). Uses Word's real
    Heading 1 style (outline level 0) so the document gets a working table of contents /
    navigation pane, not just bold text that looks like a heading.
    """
    heading = doc.add_heading(level=1)
    run = heading.add_run(text)
    run.font.name = brand.FONT_HEADING
    run.font.size = brand.SIZE_SECTION_HEADING
    run.font.bold = True
    run.font.color.rgb = brand.MIDNIGHT


def add_subheading(doc: Document, text: str) -> None:
    """Subsection heading (e.g. "3.1 Business income change"). Real Heading 2 style."""
    heading = doc.add_heading(level=2)
    run = heading.add_run(text)
    run.font.name = brand.FONT_HEADING
    run.font.size = brand.SIZE_SUBHEADING
    run.font.bold = True
    run.font.color.rgb = brand.MIDNIGHT


def add_body_paragraph(doc: Document, text: str, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = brand.LINE_SPACING_BODY
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = brand.FONT_BODY
    run.font.size = brand.SIZE_BODY
    run.font.color.rgb = brand.GREY_700
    run.font.italic = italic


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = brand.FONT_BODY
    run.font.size = brand.SIZE_CAPTION
    run.font.color.rgb = brand.GREY_700
    run.font.italic = True


def add_dashboard_visual(doc: Document, visual) -> None:
    """Embeds the real screenshot if dashboard_visuals found one, otherwise prints the
    template's own placeholder text verbatim (matching how an unfilled template slot reads).
    `visual` is a dashboard_visuals.lookup.DashboardVisual.
    """
    if visual.found:
        doc.add_picture(str(visual.path), width=Inches(6.0))
    else:
        add_caption(doc, visual.placeholder_text)


def add_table(doc: Document, headers: list, rows: list) -> None:
    """Orange header row (white bold text), alternating white/Field-50 body rows, thin grey
    borders, DXA (not percentage) column widths -- skill's "Tables" section.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _apply_table_borders(table)

    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, headers):
        _shade_cell(cell, brand.TABLE_HEADER_BG)
        _set_cell_text(cell, text, color=brand.WHITE, bold=True)

    for i, row_values in enumerate(rows):
        cells = table.add_row().cells
        bg = brand.TABLE_ALT_ROW_BG if i % 2 == 1 else None
        for cell, value in zip(cells, row_values):
            if bg:
                _shade_cell(cell, bg)
            _set_cell_text(cell, str(value), color=brand.GREY_700, bold=False)
    return table


def _set_cell_text(cell, text: str, color, bold: bool) -> None:
    # Deliberately NOT `cell.text = ""` first -- a freshly added table cell's first paragraph
    # already has zero runs, and setting cell.text leaves a stray empty run behind that
    # `cell.paragraphs[0].runs[0]` would then return instead of the real, formatted one
    # (confirmed by reading back a generated cell: runs[0] was '' with no font.color at all,
    # runs[1] was the actual text -- caught by test_table_header_row_is_orange_with_white_bold_text).
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = brand.FONT_BODY
    run.font.size = brand.SIZE_BODY
    run.font.color.rgb = color
    run.font.bold = bold
    for margin, value in (("top", 80), ("bottom", 80), ("left", 120), ("right", 120)):
        _set_cell_margin(cell, margin, value)


def _set_cell_margin(cell, side: str, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    node = OxmlElement(f"w:{side}")
    node.set(qn("w:w"), str(twips))
    node.set(qn("w:type"), "dxa")
    mar.append(node)


def _shade_cell(cell, hex_color: str) -> None:
    """ShadingType.CLEAR equivalent -- the skill explicitly warns against SOLID shading."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _apply_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")  # thin
        node.set(qn("w:color"), brand.TABLE_BORDER_GREY)
        borders.append(node)
    tbl_pr.append(borders)
