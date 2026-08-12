"""Verifies the renderer's own output actually matches wvi-docx.skill's spec -- reading back
the generated OOXML rather than trusting that the helper code did what it says, since a typo
in a hex string or a units conversion would otherwise pass silently until someone opened the
file and looked at it.
"""

from docx import Document
from docx.oxml.ns import qn

from report_render import brand, docx_helpers as h


def _new_doc() -> Document:
    return Document()


def test_setup_page_is_us_letter_with_one_inch_margins():
    doc = _new_doc()
    h.setup_page(doc)
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11.0
    assert round(section.left_margin.inches, 2) == 1.0
    assert round(section.top_margin.inches, 2) == 1.0


def test_title_uses_midnight_bold_inter():
    doc = _new_doc()
    h.add_title(doc, "Test Title")
    run = doc.paragraphs[-1].runs[0]
    assert run.font.color.rgb == brand.MIDNIGHT
    assert run.font.bold is True
    assert run.font.name == "Inter"


def test_body_paragraph_uses_grey_700():
    doc = _new_doc()
    h.add_body_paragraph(doc, "Some body text.")
    run = doc.paragraphs[-1].runs[0]
    assert run.font.color.rgb == brand.GREY_700
    assert run.font.size == brand.SIZE_BODY


def test_section_heading_is_real_heading_1_style():
    doc = _new_doc()
    h.add_section_heading(doc, "Part 1 -- Financial Access")
    paragraph = doc.paragraphs[-1]
    assert paragraph.style.name == "Heading 1"
    run = paragraph.runs[0]
    assert run.font.color.rgb == brand.MIDNIGHT
    assert run.font.bold is True


def test_table_header_row_is_orange_with_white_bold_text():
    doc = _new_doc()
    table = h.add_table(doc, ["A", "B"], [["1", "2"]])
    header_cell = table.rows[0].cells[0]
    shd = header_cell._tc.find(f".//{qn('w:shd')}")
    assert shd.get(qn("w:fill")).upper() == brand.TABLE_HEADER_BG
    run = header_cell.paragraphs[0].runs[0]
    assert run.font.color.rgb == brand.WHITE
    assert run.font.bold is True


def test_table_never_uses_solid_shading():
    # The skill explicitly warns: "Always use ShadingType.CLEAR -- never SOLID."
    doc = _new_doc()
    table = h.add_table(doc, ["A"], [["1"], ["2"]])
    for row in table.rows:
        shd = row.cells[0]._tc.find(f".//{qn('w:shd')}")
        if shd is not None:
            assert shd.get(qn("w:val")) == "clear"


def test_table_alternating_rows_use_field_50():
    doc = _new_doc()
    table = h.add_table(doc, ["A"], [["row0"], ["row1"], ["row2"]])
    # row index 0 (first data row) unshaded, row index 1 shaded Field 50, per the skill's
    # alternating-row rule -- header row (table.rows[0]) is separate from this alternation.
    row0_shd = table.rows[1].cells[0]._tc.find(f".//{qn('w:shd')}")
    row1_shd = table.rows[2].cells[0]._tc.find(f".//{qn('w:shd')}")
    assert row0_shd is None
    assert row1_shd is not None
    assert row1_shd.get(qn("w:fill")).upper() == brand.TABLE_ALT_ROW_BG


def test_divider_is_a_paragraph_border_not_a_table():
    doc = _new_doc()
    tables_before = len(doc.tables)
    h.add_divider(doc)
    assert len(doc.tables) == tables_before  # no table was added
    p_pr = doc.paragraphs[-1]._p.find(qn("w:pPr"))
    borders = p_pr.find(qn("w:pBdr"))
    bottom = borders.find(qn("w:bottom"))
    assert bottom.get(qn("w:color")).upper() == "FF5515"


def test_header_embeds_a_real_image_not_text():
    doc = _new_doc()
    h.setup_header(doc)
    header_paragraph = doc.sections[0].header.paragraphs[0]
    blips = header_paragraph.runs[-1]._element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    )
    assert len(blips) == 1


def test_footer_contains_a_real_page_field_not_a_hardcoded_number():
    doc = _new_doc()
    h.setup_footer(doc)
    footer_paragraph = doc.sections[0].footer.paragraphs[0]
    instr_texts = footer_paragraph._p.findall(f".//{qn('w:instrText')}")
    assert any(node.text == "PAGE" for node in instr_texts)


def test_footer_with_no_run_id_or_model_matches_original_plain_page_number():
    doc = _new_doc()
    h.setup_footer(doc)
    footer_paragraph = doc.sections[0].footer.paragraphs[0]
    assert footer_paragraph.runs[0].text == ""


def test_footer_includes_run_id_and_model_version_for_traceability():
    # Added after a real incident: the same wave's qualitative tagging produced materially
    # different findings across two runs, with nothing in the document saying which run's
    # numbers a reader was looking at.
    doc = _new_doc()
    h.setup_footer(doc, run_id="Test4", model_version="claude-sonnet-5")
    footer_paragraph = doc.sections[0].footer.paragraphs[0]
    assert "Run Test4" in footer_paragraph.runs[0].text
    assert "claude-sonnet-5" in footer_paragraph.runs[0].text


def test_footer_handles_run_id_only():
    doc = _new_doc()
    h.setup_footer(doc, run_id="Test4")
    footer_paragraph = doc.sections[0].footer.paragraphs[0]
    assert "Run Test4" in footer_paragraph.runs[0].text
