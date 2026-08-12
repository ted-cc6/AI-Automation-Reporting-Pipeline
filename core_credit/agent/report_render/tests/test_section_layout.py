"""Smoke-tests render_report() against a full synthetic CoreCreditImpactReport -- every field
every render_* function touches, filled with minimal-but-valid data. This is what caught the
real metric_label/GapComparison bug during development (render_child_wellbeing assumed a field
that doesn't exist on GapComparison) before it ever reached real data. Doesn't hit the real CSV,
real LLM, or real driver output -- pure structural coverage, fast enough to run every time.
"""

from docx import Document

from schemas.agency import AgencySection
from schemas.business_household_impact import BusinessHouseholdImpactSection
from schemas.child_wellbeing import ChildWellbeingSection
from schemas.client_profile import ClientProfileSection
from schemas.client_protection import ClientProtectionSection
from schemas.client_satisfaction import ClientSatisfactionSection, NPSResult
from schemas.client_voices import ClientVoicesSection
from schemas.common import (
    BenchmarkComparison,
    GapComparison,
    MetricResult,
    QualitativeSynthesis,
    RankedOption,
    RankedOptions,
    SegmentAxis,
    SegmentedValue,
    ThemeFinding,
    Verbatim,
    WrittenText,
)
from schemas.executive_summary import ExecutiveSummarySection, ThemeScore
from schemas.financial_access import FinancialAccessSection
from schemas.gender_scorecard import GenderScorecardRow, GenderScorecardSection
from schemas.poverty_likelihood import PovertyLikelihoodSection
from schemas.report import CoreCreditImpactReport
from schemas.resilience import ResilienceSection

from report_render.section_layout import render_executive_summary, render_report


def _wt(subsection_id: str) -> WrittenText:
    return WrittenText(subsection_id=subsection_id, text="Sample analysis text.", word_count=3, within_cap=True, ungrounded_percentages=[])


def _metric(metric_id: str, share: float = 0.5, n: int = 100) -> MetricResult:
    return MetricResult(
        metric_id=metric_id, label=metric_id,
        overall=SegmentedValue(axis=SegmentAxis.OVERALL, value_label="Overall", share=share, n=n),
    )


def _verbatim(country: str = "KEN") -> Verbatim:
    return Verbatim(quote="Sample quote.", gender="Female", age=30, country=country, source_field="test")


def _ranked() -> RankedOptions:
    return RankedOptions(base_n=100, options=[RankedOption(label="Option A", share=0.6, n=60)])


def _qualitative() -> QualitativeSynthesis:
    return QualitativeSynthesis(
        source_field="test", base_n=10,
        themes=[ThemeFinding(theme="Sample theme", frequency=5, share_of_respondents=0.5, representative_verbatims=[_verbatim()])],
    )


def _gap() -> GapComparison:
    return GapComparison(group_a_label="Caregiver", group_a_share=0.6, group_a_n=50, group_b_label="Non-caregiver", group_b_share=0.5, group_b_n=50, gap=0.1)


def _full_report() -> CoreCreditImpactReport:
    client_profile = ClientProfileSection(
        n_respondents=100, n_mfis=5, n_countries=5,
        gender_split=_ranked(), age=_metric("age"), household_size=_metric("household_size"),
        loan_cycle_mix=_ranked(), household_head_status=_ranked(), education_level=_ranked(),
        main_income_source=_ranked(), populated_segments=[], unavailable_segments=[],
        analysis_text=_wt("client-profile"),
    )
    financial_access = FinancialAccessSection(
        first_time_access=_metric("first_time_access"), first_time_access_analysis=_wt("1.1"),
        alternative_lender_hard_to_find=_metric("alternative_lender_hard_to_find"), alternative_lender_hard_to_find_analysis=_wt("1.2"),
        insight_text=_wt("1-insight"), insight_verbatims=[_verbatim()],
    )
    poverty_likelihood = PovertyLikelihoodSection(
        country_results=[], poverty_line_shares=[_metric("poverty_likelihood_USD190day2011PPP")],
        poverty_line_shares_analysis=_wt("2.1"), national_comparison=[], national_comparison_analysis=_wt("2.2"),
        na_footnote="Some countries excluded.", insight_text=_wt("2-insight"), insight_verbatims=[],
    )
    business_household_impact = BusinessHouseholdImpactSection(
        business_income_change=_metric("business_income_change"), business_income_analysis=_wt("3.1"),
        quality_of_life_change=_metric("quality_of_life_change"), quality_of_life_analysis=_wt("3.2"),
        qol_drivers=_qualitative(), insight_text=_wt("3-insight"), insight_verbatims=[_verbatim()],
    )
    child_wellbeing = ChildWellbeingSection(
        improved_child_wellbeing=_metric("improved_child_wellbeing"), what_improved=_ranked(),
        other_improvements_qualitative=_qualitative(), improved_child_wellbeing_analysis=_wt("4.1"),
        caregiver_vs_other=[_gap() for _ in range(8)],  # must match CAREGIVER_TABLE_LABELS length
        caregiver_vs_other_analysis=_wt("4.2"), insight_text=_wt("4-insight"), insight_verbatims=[_verbatim()],
    )
    client_protection = ClientProtectionSection(
        financial_worry_decreased=_metric("financial_worry_decreased"), financial_worry_decreased_analysis=_wt("5.1"),
        loan_terms_clear=_metric("loan_terms_clear"), loan_terms_clear_analysis=_wt("5.2"),
        complaints_mechanism_trusted=_metric("complaints_mechanism_trusted"), complaints_mechanism_trusted_analysis=_wt("5.3"),
        no_unfair_treatment=_metric("no_unfair_treatment"), no_unfair_treatment_analysis=_wt("5.4"),
        did_not_reduce_food=_metric("did_not_reduce_food"), did_not_reduce_food_analysis=_wt("5.5"),
        protection_signals=_qualitative(), insight_text=_wt("5-insight"), insight_verbatims=[],
    )
    agency = AgencySection(
        loan_purpose_achieved_fully=_metric("loan_purpose_achieved_fully"), loan_purpose_achieved_partially=_metric("loan_purpose_achieved_partially"),
        loan_purpose_achieved_analysis=_wt("6.1"), household_influence_improved=_metric("household_influence_improved"),
        household_influence_improved_analysis=_wt("6.2"), community_respect_improved=_metric("community_respect_improved"),
        community_respect_improved_analysis=_wt("6.3"), insight_text=_wt("6-insight"), insight_verbatims=[_verbatim()],
    )
    resilience = ResilienceSection(
        savings_increased=_metric("savings_increased"), savings_increased_analysis=_wt("7.1"),
        shock_incidence=_metric("shock_incidence"), shock_impacts=_ranked(), shock_incidence_analysis=_wt("7.2"),
        coping_mechanisms=_ranked(), negative_coping_share=_metric("negative_coping_share"),
        other_coping_qualitative=_qualitative(), coping_mechanisms_analysis=_wt("7.3"),
        vf_reduced_shock_severity=_metric("vf_reduced_shock_severity"), vf_reduced_shock_severity_analysis=_wt("7.4"),
        insight_text=_wt("7-insight"), insight_verbatims=[_verbatim()],
    )
    client_satisfaction = ClientSatisfactionSection(
        nps=NPSResult(score=50, promoter_share=0.6, passive_share=0.2, detractor_share=0.2, n=100),
        nps_analysis=_wt("8.1"), promoter_drivers=_ranked(), detractor_pain_points=_ranked(),
        nps_followup_themes=[_qualitative(), _qualitative(), _qualitative()],
        drivers_analysis=_wt("8.2"), insight_text=_wt("8-insight"), insight_verbatims=[_verbatim()],
    )
    executive_summary = ExecutiveSummarySection(
        theme_scores=[ThemeScore(theme_name="Financial Access", metric_label="First-time access", headline_value=0.4)],
        n_respondents=100, n_mfis=5, n_countries=5, reporting_period="2026 Q2", generated_date="2026-08-06",
        analysis_text=_wt("executive-summary"),
    )
    gender_scorecard = GenderScorecardSection(
        rows=[GenderScorecardRow(metric_label="First-time access", female_share=0.4, male_share=0.5, gap=-0.1)],
        analysis_text=_wt("gender-scorecard"), insight_text=_wt("gender-insight"), insight_verbatims=[_verbatim()],
    )
    client_voices = ClientVoicesSection(green_lights=[_verbatim()], red_flags=[_verbatim()])

    return CoreCreditImpactReport(
        reporting_period="2026 Q2", generated_at="2026-08-06",
        client_profile=client_profile, financial_access=financial_access, poverty_likelihood=poverty_likelihood,
        business_household_impact=business_household_impact, child_wellbeing=child_wellbeing,
        client_protection=client_protection, agency=agency, resilience=resilience,
        client_satisfaction=client_satisfaction, executive_summary=executive_summary,
        gender_scorecard=gender_scorecard, client_voices=client_voices,
    )


def test_render_report_does_not_crash_on_a_full_synthetic_report():
    doc = Document()
    render_report(doc, _full_report())
    assert len(doc.paragraphs) > 0


def test_render_report_produces_every_expected_top_level_heading():
    doc = Document()
    render_report(doc, _full_report())
    headings = {p.text for p in doc.paragraphs if p.style.name == "Heading 1"}
    for expected in [
        "Client Profile & Methodology", "Executive Summary", "Part 1 -- Financial Access",
        "Part 2 -- Poverty Likelihood (PPI)", "Part 3 -- Business & Household Impact",
        "Part 4 -- Child Wellbeing", "Part 5 -- Client Protection", "Part 6 -- Agency",
        "Part 7 -- Resilience", "Part 8 -- Client Satisfaction", "Part 9 -- Gender",
        "Part 10 -- Client Voices",
    ]:
        assert expected in headings, f"missing heading: {expected}"


def test_render_report_produces_exactly_three_tables():
    doc = Document()
    render_report(doc, _full_report())
    assert len(doc.tables) == 3  # Executive Summary, Child Wellbeing 4.2, Gender scorecard


def test_executive_summary_table_shows_comparable_value_not_just_headline():
    # Regression test for a real incident: the table showed Resilience's loose headline
    # (77.5%) next to the 16.0% MFI Index benchmark, while the prose below correctly used the
    # matched-basis figure (27.8%) for the same comparison -- two different numbers for the
    # same claim, with nothing explaining why. The table must expose the comparable figure too.
    section = ExecutiveSummarySection(
        theme_scores=[
            ThemeScore(
                theme_name="Resilience",
                metric_label="Savings increased",
                headline_value=0.775,
                benchmark=BenchmarkComparison(external_mfi_index=0.16, external_mfi_index_year=2025),
                benchmark_comparable_value=0.278,
            )
        ],
        n_respondents=100,
        n_mfis=1,
        n_countries=1,
        reporting_period="Q1 2026",
        generated_date="2026-01-01",
    )
    doc = Document()
    render_executive_summary(doc, section)
    table = doc.tables[0]
    header_row = [c.text for c in table.rows[0].cells]
    data_row = [c.text for c in table.rows[1].cells]
    assert "Comparable to Benchmark" in header_row
    assert "77.5%" in data_row  # headline still shown
    assert "27.8%" in data_row  # comparable figure now also shown, not silently dropped
    assert "16.0%" in data_row  # the actual benchmark


def test_render_report_works_when_cross_cutting_sections_are_none():
    report = _full_report()
    report.executive_summary = None
    report.gender_scorecard = None
    report.client_voices = None
    doc = Document()
    render_report(doc, report)  # must not crash
    headings = {p.text for p in doc.paragraphs if p.style.name == "Heading 1"}
    assert "Executive Summary" not in headings
    assert "Part 9 -- Gender" not in headings
    assert "Part 10 -- Client Voices" not in headings
