from docx import Document

from schemas.common import QualitativeSynthesis, ThemeFinding, Verbatim
from report_render.section_layout import add_protection_signals


def _theme(label: str, freq: int, severity=None) -> ThemeFinding:
    return ThemeFinding(
        theme=label, frequency=freq, severity=severity,
        representative_verbatims=[Verbatim(quote="Sample quote.", country="KEN", source_field="test")],
    )


def _doc_text(doc) -> list:
    return [p.text for p in doc.paragraphs]


def test_no_qualitative_shows_honest_placeholder():
    doc = Document()
    add_protection_signals(doc, None)
    text = " ".join(_doc_text(doc))
    assert "No client-protection signals" in text


def test_empty_themes_list_shows_honest_placeholder():
    doc = Document()
    add_protection_signals(doc, QualitativeSynthesis(source_field="x", base_n=100, themes=[]))
    text = " ".join(_doc_text(doc))
    assert "No client-protection signals" in text


def test_groups_themes_by_severity_tier_in_high_medium_low_order():
    qual = QualitativeSynthesis(
        source_field="x", base_n=100,
        themes=[_theme("Minor complaint", 5, "low"), _theme("Coercion", 3, "high"), _theme("Pressure tactics", 4, "medium")],
    )
    doc = Document()
    add_protection_signals(doc, qual)
    text = _doc_text(doc)
    high_idx = text.index("High severity")
    medium_idx = text.index("Medium severity")
    low_idx = text.index("Low severity")
    assert high_idx < medium_idx < low_idx
    assert any("Coercion" in t for t in text[high_idx:medium_idx])
    assert any("Pressure tactics" in t for t in text[medium_idx:low_idx])
    assert any("Minor complaint" in t for t in text[low_idx:])


def test_caps_each_tier_to_top_n_by_frequency():
    themes = [_theme(f"High theme {i}", freq, "high") for i, freq in enumerate([50, 40, 30, 20, 10, 5])]
    qual = QualitativeSynthesis(source_field="x", base_n=1000, themes=themes)
    doc = Document()
    add_protection_signals(doc, qual, cap_per_tier=3)
    text = " ".join(_doc_text(doc))
    assert "High theme 0" in text
    assert "High theme 1" in text
    assert "High theme 2" in text
    assert "High theme 3" not in text  # 4th-ranked, beyond the cap
    assert "High theme 5" not in text


def test_theme_with_no_severity_is_grouped_under_other_not_dropped():
    qual = QualitativeSynthesis(source_field="x", base_n=10, themes=[_theme("Unclassified", 2, severity=None)])
    doc = Document()
    add_protection_signals(doc, qual)
    text = " ".join(_doc_text(doc))
    assert "Unclassified" in text
    assert "Other" in text
