from unittest.mock import patch

from dashboard_visuals.lookup import DashboardVisual
from docx import Document

from report_render.section_layout import render_report
from report_render.tests.test_section_layout import _full_report


def test_render_report_uses_prefetched_visuals_instead_of_hitting_disk():
    prefetched = {
        sid: DashboardVisual(subsection_id=sid, found=False, path=None, placeholder_text=f"PREFETCHED-{sid}")
        for sid in ["client-profile", "1.1", "1.2", "2.1", "2.2", "3.1", "3.2", "4.1", "4.2",
                    "5.1", "5.2", "5.3", "5.4", "5.5", "6.1", "6.2", "6.3", "7.1", "7.2", "7.3",
                    "7.4", "8.1", "8.2", "executive-summary", "gender-scorecard", "client-voices"]
    }
    doc = Document()
    with patch("report_render.section_layout.find_dashboard_visual") as mock_find:
        render_report(doc, _full_report(), dashboard_visuals=prefetched)
    mock_find.assert_not_called()
    text = " ".join(p.text for p in doc.paragraphs)
    assert "PREFETCHED-1.1" in text


def test_render_report_falls_back_to_live_lookup_when_not_given_prefetched_visuals():
    doc = Document()
    with patch("report_render.section_layout.find_dashboard_visual") as mock_find:
        mock_find.return_value = DashboardVisual(subsection_id="x", found=False, path=None, placeholder_text="LIVE")
        render_report(doc, _full_report())
    assert mock_find.call_count == 26
