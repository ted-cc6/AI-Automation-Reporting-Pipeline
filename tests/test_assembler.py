"""
Unit tests for generation/assembler.py's scope-aware cover page (Phase 5).
The core guarantee under test: an unscoped (global-portfolio) run must
produce a byte-identical cover heading/subtitle to what assembler.py always
produced, so adding single-country support can't change the global report.
Run: pytest tests/test_assembler.py -v
"""
from __future__ import annotations

import json

import pytest
from docx import Document

import generation.assembler as assembler
from generation.assembler import (
    _add_drivers_table,
    _add_executive_summary,
    _add_insight_box,
    _add_protection_signals_annex,
    _add_protection_signals_summary,
    _group_header,
    _load_analysis_meta,
    assemble,
    build_part_6,
    build_part_7,
    build_part_10,
)


def _make_run(tmp_path, run_id: str, meta: dict) -> None:
    run_dir = tmp_path / "runs" / run_id
    run_dir.mkdir(parents=True)
    (run_dir / "analysis_results.json").write_text(json.dumps({"meta": meta}), encoding="utf-8")


def _make_run_with_parts(tmp_path, run_id: str, meta: dict, parts: dict, qual: "dict | None" = None) -> None:
    run_dir = tmp_path / "runs" / run_id
    run_dir.mkdir(parents=True)
    (run_dir / "analysis_results.json").write_text(
        json.dumps({"meta": meta, "parts": parts}), encoding="utf-8"
    )
    if qual is not None:
        (run_dir / "qualitative_results.json").write_text(json.dumps(qual), encoding="utf-8")


# ---------------------------------------------------------------------------
# _load_analysis_meta
# ---------------------------------------------------------------------------

class TestLoadAnalysisMeta:
    def test_missing_run_dir_returns_empty_dict(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        assert _load_analysis_meta("no_such_run") == {}

    def test_valid_file_returns_meta_block(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "good_run", {"country": "vietnam", "n_total": 154})
        assert _load_analysis_meta("good_run") == {"country": "vietnam", "n_total": 154}


# ---------------------------------------------------------------------------
# _group_header / build_part_6 (R-011)
# ---------------------------------------------------------------------------

class TestGroupHeader:
    def test_no_qualifier_matches_prior_behaviour(self):
        groups = {"female": {"label": "Female", "n": 900}}
        assert _group_header(groups, "female", "Female") == "Female (n=900)"

    def test_qualifier_folds_into_same_parenthetical(self):
        groups = {"claimant": {"label": "Claimant", "n": 55, "qualifier": "filed"}}
        assert _group_header(groups, "claimant", "Claimant") == "Claimant (filed, n=55)"

    def test_no_n_returns_bare_label_even_with_qualifier(self):
        groups = {"claimant": {"label": "Claimant", "qualifier": "filed"}}
        assert _group_header(groups, "claimant", "Claimant") == "Claimant"

    def test_missing_key_uses_fallback_label(self):
        assert _group_header({}, "claimant", "Claimant (filed)") == "Claimant (filed)"


class TestBuildPart6:
    def _package(self, **overrides) -> dict:
        pkg = {
            "title": "Claimant vs. Did-Not-File Outcomes",
            "sections": {"insight": {"verbatims": []}},
            "visuals": [],
            "scorecard": [{
                "label": "Coverage Understanding",
                "group_a_value": "85.5%",
                "group_b_value": "70.1%",
                "significant": True,
                "population": None,
                "sig_test_note": None,
            }],
            "groups": {
                "claimant": {"label": "Claimant", "n": 55, "qualifier": "filed"},
                "non_claimant": {"label": "Did not file", "n": 69},
            },
        }
        pkg.update(overrides)
        return pkg

    def _paragraph_texts(self, doc) -> list:
        return [p.text for p in doc.paragraphs]

    def test_headers_state_the_population_inline(self):
        doc = Document()
        build_part_6(doc, self._package(), {"narrative": ""})
        table = doc.tables[0]
        header_cells = [c.text for c in table.rows[0].cells]
        assert header_cells == ["Metric", "Claimant (filed, n=55)", "Did not file (n=69)", "Sig.*"]

    def test_population_scope_note_renders_beneath_table(self):
        doc = Document()
        build_part_6(doc, self._package(), {"narrative": ""})
        texts = "\n".join(self._paragraph_texts(doc))
        assert "restricted to clients who reported an insured event" in texts

    def test_no_retired_labels_anywhere(self):
        doc = Document()
        build_part_6(doc, self._package(), {"narrative": ""})
        texts = "\n".join(self._paragraph_texts(doc)).lower()
        for cell in doc.tables[0].rows[0].cells:
            texts += "\n" + cell.text.lower()
        assert "non-filer" not in texts
        assert "non filer" not in texts
        assert "non-claimant" not in texts
        assert "non claimant" not in texts


# ---------------------------------------------------------------------------
# _add_insight_box (R-010)
# ---------------------------------------------------------------------------

def _doc_paragraph_texts(doc) -> list:
    return [p.text for p in doc.paragraphs]


class TestAddInsightBox:
    def test_no_verbatims_renders_nothing(self):
        doc = Document()
        _add_insight_box(doc, "Some narrated absence of data.", [])
        assert _doc_paragraph_texts(doc) == []

    def test_verbatims_with_no_text_key_renders_nothing(self):
        doc = Document()
        verbatims = [{"profile": {"sex": "female"}}, {"text": "", "profile": {}}]
        _add_insight_box(doc, "Insight narrative.", verbatims)
        assert _doc_paragraph_texts(doc) == []

    def test_heading_never_renders_without_a_verbatim(self):
        doc = Document()
        _add_insight_box(doc, "Insight narrative.", [])
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Key Qualitative Insights" not in headings

    def test_insight_text_never_renders_without_a_verbatim(self):
        # The exact regression this fix targets: insight_text generated by
        # the same LLM call that had nothing to quote can itself narrate
        # the absence (C-014) -- omitting the whole block removes it.
        doc = Document()
        _add_insight_box(doc, "This data has not yet been analyzed for this report.", [])
        assert "has not yet been analyzed" not in "\n".join(_doc_paragraph_texts(doc))

    def test_one_valid_verbatim_renders_heading_and_quote(self):
        doc = Document()
        verbatims = [{"text": "It changed my life.", "profile": {"sex": "female", "age": 34}}]
        _add_insight_box(doc, "Insight narrative.", verbatims)
        texts = _doc_paragraph_texts(doc)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Key Qualitative Insights" in headings
        assert any('"It changed my life."' == t for t in texts)

    def test_mixed_valid_and_empty_verbatims_only_renders_valid_ones(self):
        doc = Document()
        verbatims = [
            {"text": "", "profile": {}},
            {"text": "Real quote here.", "profile": {}},
        ]
        _add_insight_box(doc, "", verbatims)
        texts = _doc_paragraph_texts(doc)
        quotes = [t for t in texts if t.startswith('"')]
        assert quotes == ['"Real quote here."']

    def test_more_than_three_verbatims_still_caps_at_three(self):
        doc = Document()
        verbatims = [{"text": f"Quote {i}", "profile": {}} for i in range(5)]
        _add_insight_box(doc, "", verbatims)
        texts = _doc_paragraph_texts(doc)
        quotes = [t for t in texts if t.startswith('"')]
        assert len(quotes) == 3


# ---------------------------------------------------------------------------
# assemble() cover page
# ---------------------------------------------------------------------------

class TestAssembleCoverPage:
    def _cover_paragraphs(self, doc_path) -> list[str]:
        doc = Document(str(doc_path))
        return [p.text for p in doc.paragraphs[:4]]

    def test_no_meta_matches_original_hardcoded_global_cover(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "no_meta_run", {})
        out = tmp_path / "out.docx"
        assemble([], {}, "no_meta_run", out)
        period = assembler.format_period_label("no_meta_run")
        paras = self._cover_paragraphs(out)
        assert paras[0] == "VisionFund International"
        assert paras[1] == f"Insurance Impact Report: Global Portfolio, {period}"

    def test_default_country_matches_original_hardcoded_global_cover(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "default_run", {"country": "default", "country_label": "Default", "n_total": 2111})
        out = tmp_path / "out.docx"
        assemble([], {}, "default_run", out)
        paras = self._cover_paragraphs(out)
        assert "Global Portfolio" in paras[1]
        assert paras[2] == "Covering 2,111 client responses across the VisionFund insurance portfolio."

    def test_scoped_country_uses_country_label_in_title_and_subtitle(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "vietnam_run", {"country": "vietnam", "country_label": "Vietnam", "n_total": 154})
        out = tmp_path / "out.docx"
        assemble([], {}, "vietnam_run", out)
        paras = self._cover_paragraphs(out)
        period = assembler.format_period_label("vietnam_run")
        assert paras[1] == f"Insurance Impact Report: Vietnam, {period}"
        assert paras[2] == "Covering 154 client responses from Vietnam."
        assert "Global Portfolio" not in paras[1]

    def test_scoped_country_without_label_falls_back_to_titlecase(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "kenya_run", {"country": "kenya", "n_total": 300})
        out = tmp_path / "out.docx"
        assemble([], {}, "kenya_run", out)
        paras = self._cover_paragraphs(out)
        assert "Kenya" in paras[1]

    def test_larco_rollup_uses_lacro_regional_title_not_global(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "larco_run", {"country": "default", "dataset_schema": "larco", "n_total": 1355})
        out = tmp_path / "out.docx"
        assemble([], {}, "larco_run", out)
        paras = self._cover_paragraphs(out)
        assert "LACRO Regional Portfolio" in paras[1]
        assert "Global Portfolio" not in paras[1]
        assert paras[2] == (
            "Covering 1,355 client responses across VisionFund's LACRO "
            "(Latin America and Caribbean Regional Office) insurance portfolio."
        )
        # The internal dataset_schema value is spelled "larco"; nothing
        # reader-facing may say "LARCO" -- see report_scopes.py.
        assert "LARCO" not in paras[1]
        assert "LARCO" not in paras[2]

    def test_report_scope_lacro_uses_lacro_regional_title_not_global(self, tmp_path, monkeypatch):
        # The real bug this guards against: a report_scope=="lacro" run on
        # the unified schema (country="default") previously fell through to
        # "Global Portfolio" since only dataset_schema=="larco" was checked.
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "lacro_scope_run", {
            "country": "default", "dataset_schema": "africa_vietnam",
            "report_scope": "lacro",
            "report_scope_label": "LACRO (Latin America and Caribbean Regional Office)",
            "n_total": 1721,
        })
        out = tmp_path / "out.docx"
        assemble([], {}, "lacro_scope_run", out)
        paras = self._cover_paragraphs(out)
        assert "LACRO Regional Portfolio" in paras[1]
        assert "Global Portfolio" not in paras[1]
        assert paras[2] == (
            "Covering 1,721 client responses across VisionFund's LACRO "
            "(Latin America and Caribbean Regional Office) insurance portfolio."
        )
        # A separate, later bug (round 3): this branch hardcoded "LARCO"
        # directly into the title/subtitle instead of pulling the correct
        # reader-facing spelling from report_scopes.py.
        assert "LARCO" not in paras[1]
        assert "LARCO" not in paras[2]

    def test_report_scope_africa_uses_its_own_label_in_title_and_subtitle(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "africa_scope_run", {
            "country": "default", "dataset_schema": "africa_vietnam",
            "report_scope": "africa", "report_scope_label": "Africa and Asia",
            "n_total": 2091,
        })
        out = tmp_path / "out.docx"
        assemble([], {}, "africa_scope_run", out)
        paras = self._cover_paragraphs(out)
        assert paras[1] == f"Insurance Impact Report: Africa and Asia Portfolio, {assembler.format_period_label('africa_scope_run')}"
        assert paras[2] == "Covering 2,091 client responses across VisionFund's Africa and Asia insurance portfolio."
        assert "Global Portfolio" not in paras[1]

    def test_larco_single_country_still_uses_country_label(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "ecuador_run", {"country": "ecuador", "country_label": "Ecuador", "dataset_schema": "larco", "n_total": 400})
        out = tmp_path / "out.docx"
        assemble([], {}, "ecuador_run", out)
        paras = self._cover_paragraphs(out)
        assert paras[1] == f"Insurance Impact Report: Ecuador, {assembler.format_period_label('ecuador_run')}"
        assert "LARCO" not in paras[1]

    def test_missing_n_total_omits_subtitle_paragraph_in_both_scopes(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "no_n_total_run", {"country": "default"})
        out = tmp_path / "out.docx"
        assemble([], {}, "no_n_total_run", out)
        paras = self._cover_paragraphs(out)
        # Next paragraph after the heading should jump straight to "Generated:"
        assert paras[2].startswith("Generated:")


# ---------------------------------------------------------------------------
# _add_drivers_table -- the drivers-table not_applicable marker (the gap
# picked up after Phase 5: a population-exclusive driver's table row must
# show "NOT APPLICABLE", not "SUPPRESSED", for a human reader too).
# ---------------------------------------------------------------------------

class TestAddDriversTable:
    def _table_rows(self, driver_row: dict) -> list:
        from docx import Document
        doc = Document()
        _add_drivers_table(doc, {"drivers_data": [driver_row], "drivers_table": {}})
        table = doc.tables[-1]
        return [[cell.text for cell in row.cells] for row in table.rows]

    def test_not_applicable_driver_renders_marker_row(self):
        rows = self._table_rows({
            "label": "Renewal Intent", "rho": None, "p_value": None, "n_valid": None,
            "suppressed": True, "not_applicable": True,
        })
        assert rows[1] == ["Renewal Intent", "NOT APPLICABLE", "NOT APPLICABLE", "NOT APPLICABLE"]

    def test_ordinary_suppressed_driver_still_shows_suppressed(self):
        rows = self._table_rows({
            "label": "Coverage Understanding", "rho": None, "p_value": None, "n_valid": None,
            "suppressed": True, "not_applicable": False,
        })
        assert rows[1] == ["Coverage Understanding", "SUPPRESSED", "SUPPRESSED", "SUPPRESSED"]

    def test_normal_driver_row_unaffected(self):
        rows = self._table_rows({
            "label": "Confidence in Payout", "rho": 0.312, "p_value": 0.001, "n_valid": 1200,
            "suppressed": False, "not_applicable": False,
        })
        assert rows[1] == ["Confidence in Payout", "+0.312", "0.0010", "1200"]

    def test_tiny_p_value_floors_instead_of_truncating_to_zero(self):
        rows = self._table_rows({
            "label": "Financial Stress", "rho": -0.387, "p_value": 2.28e-38, "n_valid": 1037,
            "suppressed": False, "not_applicable": False,
        })
        assert rows[1] == ["Financial Stress", "-0.387", "<0.0001", "1037"]

    def test_prints_appendix_pointer_not_full_methodology_text(self):
        from docx import Document
        doc = Document()
        _add_drivers_table(doc, {
            "drivers_data": [{
                "label": "NPS", "rho": 0.2, "p_value": 0.01, "n_valid": 500,
                "suppressed": False, "not_applicable": False,
            }],
            "drivers_table": {},
        })
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Appendix: Methodology Notes" in body_text
        # The old ~200-word inline explanation must not be duplicated here.
        assert "Spearman rank correlation coefficient" not in body_text

    def test_default_header_is_factor_not_driver(self):
        # R-035 (docs/report_spec.md, session-11): "Driver" invites a causal
        # reading this cross-sectional data doesn't support.
        rows = self._table_rows({
            "label": "NPS", "rho": 0.2, "p_value": 0.01, "n_valid": 500,
            "suppressed": False, "not_applicable": False,
        })
        assert rows[0][0] == "Factor"
        assert "Driver" not in rows[0]

    def test_association_note_renders_when_configured(self):
        from docx import Document
        doc = Document()
        note = "These are observed associations in cross-sectional survey data and do not establish that one factor causes or determines another."
        _add_drivers_table(doc, {
            "drivers_data": [{
                "label": "NPS", "rho": 0.2, "p_value": 0.01, "n_valid": 500,
                "suppressed": False, "not_applicable": False,
            }],
            "drivers_table": {"association_note": note},
        })
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert note in body_text

    def test_association_note_omitted_when_not_configured(self):
        # Backward compatible: existing drivers_table configs with no
        # association_note key must not render an empty/placeholder line.
        from docx import Document
        doc = Document()
        _add_drivers_table(doc, {
            "drivers_data": [{
                "label": "NPS", "rho": 0.2, "p_value": 0.01, "n_valid": 500,
                "suppressed": False, "not_applicable": False,
            }],
            "drivers_table": {},
        })
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "cross-sectional" not in body_text


# ---------------------------------------------------------------------------
# Methodology appendix -- de-duplicates the Spearman ρ/p-value/N explanation
# that previously appeared verbatim under both Part 4's and Part 5's drivers
# tables into a single document appendix.
# ---------------------------------------------------------------------------

class TestMethodologyAppendix:
    def _doc_text(self, doc_path) -> str:
        doc = Document(str(doc_path))
        return "\n".join(p.text for p in doc.paragraphs)

    def test_appendix_added_when_part_4_present(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "p4_run", {"country": "default", "n_total": 100})
        out = tmp_path / "out.docx"
        assemble([{"part": "part_4", "title": "Child Wellbeing Outcomes"}], {}, "p4_run", out)
        assert "Appendix: Methodology Notes" in self._doc_text(out)

    def test_appendix_added_when_part_5_present(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "p5_run", {"country": "default", "n_total": 100})
        out = tmp_path / "out.docx"
        assemble([{"part": "part_5", "title": "Child Wellbeing"}], {}, "p5_run", out)
        assert "Appendix: Methodology Notes" in self._doc_text(out)

    def test_appendix_omitted_when_neither_part_present(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "p1_only_run", {"country": "default", "n_total": 100})
        out = tmp_path / "out.docx"
        assemble([{"part": "part_1", "title": "Client Understanding & Value Perception"}], {}, "p1_only_run", out)
        assert "Appendix: Methodology Notes" not in self._doc_text(out)

    def test_renewal_intent_example_present_for_unscoped_report(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "unscoped_run", {"country": "default", "n_total": 100})
        out = tmp_path / "out.docx"
        assemble([{"part": "part_4", "title": "Client Voice"}], {}, "unscoped_run", out)
        text = self._doc_text(out)
        assert "Renewal Intent was asked only of Vietnam's crop-insurance clients" in text
        # The scoring-direction example (a separate constant from the
        # N-variability example above) also names Renewal Intent for an
        # unscoped report, where it's a real factor in the table.
        assert 'Renewal Intent' in text
        assert '"Definitely would renew"' in text

    def test_renewal_intent_example_omitted_for_report_scope_lacro(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "lacro_scope_run",
                  {"country": "default", "n_total": 100, "report_scope": "lacro"})
        out = tmp_path / "out.docx"
        assemble([{"part": "part_4", "title": "Client Voice"}], {}, "lacro_scope_run", out)
        text = self._doc_text(out)
        assert "Vietnam" not in text
        # The scoring-direction example must also drop Renewal Intent for a
        # LACRO report -- it never appears in a LACRO drivers table, so
        # naming it here would be exactly as false as the Vietnam example.
        assert "Renewal Intent" not in text
        assert '"Definitely would renew"' not in text

    def test_renewal_intent_example_omitted_for_legacy_larco_schema_with_no_report_scope(
        self, tmp_path, monkeypatch
    ):
        # The gap this fix closes: the legacy dataset_schema=="larco" upload
        # path never offers the report_scope picker at all (see
        # CupboardWeekApp.tsx's scopeOptionsForSchema), so a real run
        # through it has report_scope == None. Before the fix,
        # _spearman_methodology_note() only checked report_scope == "lacro"
        # and would have shown this false Vietnam/Africa example text for a
        # report with zero Vietnam or African clients in it.
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "legacy_larco_run",
                  {"country": "default", "n_total": 100, "dataset_schema": "larco", "report_scope": None})
        out = tmp_path / "out.docx"
        assemble([{"part": "part_4", "title": "Client Voice"}], {}, "legacy_larco_run", out)
        text = self._doc_text(out)
        assert "Vietnam" not in text
        assert "Renewal Intent" not in text


# ---------------------------------------------------------------------------
# Part 8 (Kling Index) is deliberately dashboard-only and never gets its own
# builder -- without an explicit note, a reader just sees the part numbering
# jump from 7 straight to 9 (or 10) with nothing in the document explaining
# why. build_part_7() always renders (every Cupboard Week report includes
# Part 7), so the note lives there rather than being conditional on which
# later parts a given run happens to include.
# ---------------------------------------------------------------------------

class TestPart8AbsenceNote:
    def test_part7_explains_part8_is_dashboard_only(self):
        doc = Document()
        build_part_7(doc, {"title": "Gender Analysis", "sections": {}}, {"narrative": "x", "insight": "y"})
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Part 8 (Kling Index)" in body_text
        assert "analytics dashboard only" in body_text


# ---------------------------------------------------------------------------
# Protection signals -- short in-body summary (Part 2) + full itemized annex
# (Phase D), replacing the old single inline rendering that put the full
# itemized list (reason text + client ref per case) directly in Part 2's
# narrative flow.
# ---------------------------------------------------------------------------

_FLAGS = [
    {"id": "row_0011", "flag_type": "unfair_claim_denial", "severity": "high",
     "reason": "Client says claim was denied without explanation.",
     "profile": {"client_id": "CI-00011", "branch": "Branch A"}},
    {"id": "row_0042", "flag_type": "staff_misconduct", "severity": "medium",
     "reason": "Client reports unresponsive branch staff.",
     "profile": {"client_id": "CI-00042", "branch": "Branch B"}},
    {"id": "row_0099", "flag_type": "staff_misconduct", "severity": "low",
     "reason": "Minor communication friction reported.",
     "profile": {"client_id": "CI-00099", "branch": "Branch C"}},
]


class TestProtectionSignalsSummary:
    def _paragraphs(self, protection_flags: list) -> list:
        from docx import Document
        doc = Document()
        _add_protection_signals_summary(doc, protection_flags, assembler._compute_severity_counts(protection_flags))
        return [p.text for p in doc.paragraphs]

    def test_no_flags_renders_nothing(self):
        assert self._paragraphs([]) == []

    def test_shows_counts_by_severity_not_full_reason_text(self):
        paras = self._paragraphs(_FLAGS)
        body = "\n".join(paras)
        assert "Client Protection Signals" in body
        assert "3 client-reported protection concerns" in body
        assert "1 high" in body and "1 medium" in body and "1 low" in body
        assert "Appendix: Client Protection Signals" in body
        # The full per-case reason text must NOT leak into the short summary.
        assert "Client says claim was denied without explanation." not in body
        assert "CI-00011" not in body

    def test_singular_wording_for_exactly_one_flag(self):
        body = "\n".join(self._paragraphs(_FLAGS[:1]))
        assert "1 client-reported protection concern was identified" in body


class TestProtectionSignalsAnnex:
    def _paragraphs(self, protection_flags: list) -> list:
        from docx import Document
        doc = Document()
        _add_protection_signals_annex(doc, protection_flags, assembler._compute_severity_counts(protection_flags))
        return [p.text for p in doc.paragraphs]

    def test_no_flags_renders_nothing(self):
        assert self._paragraphs([]) == []

    def test_full_itemized_list_grouped_by_severity_with_client_refs(self):
        body = "\n".join(self._paragraphs(_FLAGS))
        assert "Appendix: Client Protection Signals" in body
        assert "High severity" in body
        assert "Unfair claim denial: Client says claim was denied without explanation. (CI-00011, Branch A)" in body
        assert "Staff misconduct: Client reports unresponsive branch staff. (CI-00042, Branch B)" in body
        assert "Staff misconduct: Minor communication friction reported. (CI-00099, Branch C)" in body

    def test_falls_back_to_row_id_when_profile_unresolved(self):
        # A flag whose row_id no longer maps into the survey dataframe (e.g.
        # a stale re-run) must still render something traceable rather than
        # silently dropping the reference.
        flags = [{"id": "row_0007", "flag_type": "coercion", "severity": "high",
                  "reason": "Unresolved case.", "profile": {}}]
        body = "\n".join(self._paragraphs(flags))
        assert "(row_0007)" in body

    def test_severity_order_high_medium_low(self):
        paras = self._paragraphs(_FLAGS)
        headings = [p for p in paras if p.endswith("severity")]
        assert headings == ["High severity", "Medium severity", "Low severity"]

    def test_same_client_multiple_concerns_is_annotated(self):
        # R-003: a client with two genuinely distinct kept concerns (after
        # qualitative/parse_results.py's client-level dedup) must read as
        # explained, not as an unexplained repeat of the same client ref.
        flags = [
            {"id": "row_0011", "flag_type": "unfair_claim_denial", "severity": "high",
             "reason": "Claim denied without explanation.",
             "profile": {"client_id": "CI-00011", "branch": "Branch A"},
             "same_client_multiple_concerns": True},
            {"id": "row_0055", "flag_type": "staff_misconduct", "severity": "medium",
             "reason": "Separate complaint about a different visit.",
             "profile": {"client_id": "CI-00011", "branch": "Branch A"},
             "same_client_multiple_concerns": True},
        ]
        body = "\n".join(self._paragraphs(flags))
        assert "Claim denied without explanation. (CI-00011, Branch A; same client, multiple concerns)" in body
        assert "Separate complaint about a different visit. (CI-00011, Branch A; same client, multiple concerns)" in body

    def test_single_concern_client_is_not_annotated(self):
        body = "\n".join(self._paragraphs(_FLAGS))
        assert "same client, multiple concerns" not in body


class TestSeverityCountsSharedBetweenSummaryAndAnnex:
    """R-017's actual acceptance criterion: the severity counts rendered in
    the Part 2 in-body summary must equal those rendered in the appendix."""

    def test_summary_and_annex_counts_agree(self):
        from docx import Document

        severity_counts = assembler._compute_severity_counts(_FLAGS)

        summary_doc = Document()
        _add_protection_signals_summary(summary_doc, _FLAGS, severity_counts)
        summary_paras = [p.text for p in summary_doc.paragraphs]

        annex_doc = Document()
        _add_protection_signals_annex(annex_doc, _FLAGS, severity_counts)
        annex_paras = [p.text for p in annex_doc.paragraphs]

        # Count annex bullets under each "<Severity> severity" heading.
        annex_counts: dict[str, int] = {}
        current = None
        for p in annex_paras:
            if p.endswith(" severity"):
                current = p.split()[0].lower()
                annex_counts[current] = 0
            elif current and p.strip():
                annex_counts[current] += 1

        for sev, count in severity_counts.items():
            assert f"{count} {sev}" in "\n".join(summary_paras), (
                f"summary doesn't state {count} {sev}"
            )
        assert annex_counts == severity_counts, (
            f"appendix bullet counts {annex_counts} disagree with severity_counts {severity_counts}"
        )


class TestAssembleProtectionSignalsRouting:
    def _doc_text(self, doc_path) -> str:
        doc = Document(str(doc_path))
        return "\n".join(p.text for p in doc.paragraphs)

    def test_flags_extracted_from_part_2_package_and_routed_to_annex(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "flags_run", {"country": "default", "n_total": 100})
        out = tmp_path / "out.docx"
        package = {
            "part": "part_2", "title": "Claims Experience",
            "sections": {"s2_3": {"qualitative": {"protection_flags": _FLAGS}}},
        }
        assemble([package], {}, "flags_run", out)
        body = self._doc_text(out)
        assert "Appendix: Client Protection Signals" in body
        assert "Client says claim was denied without explanation. (CI-00011, Branch A)" in body
        # Part 2's own inline section shows only the short summary.
        assert "3 client-reported protection concerns" in body

    def test_no_part_2_package_omits_annex(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "no_p2_run", {"country": "default", "n_total": 100})
        out = tmp_path / "out.docx"
        assemble([{"part": "part_1", "title": "Client Understanding & Value Perception"}], {}, "no_p2_run", out)
        assert "Appendix: Client Protection Signals" not in self._doc_text(out)


# ---------------------------------------------------------------------------
# Executive Summary -- headline numbers (deterministic) + top findings/
# actions + narrative (from qualitative_results.json's Task 7) + a
# consolidated data-availability caveat box (Phase D).
# ---------------------------------------------------------------------------

_EXEC_SUMMARY_PARTS = {
    "part_4": {
        "nps": {"result": {"value": 20.0, "n_valid": 500, "suppressed": False, "not_applicable": False}},
        "child_wellbeing": {"headline": {"value": 0.35, "n_valid": 400, "suppressed": False, "not_applicable": False}},
    },
    "part_3": {
        "metrics": {
            "no_prior_access": {"headline": {"value": 0.8, "n_valid": 500, "suppressed": False, "not_applicable": False}},
            "confidence_pay": {"headline": {"value": None, "n_valid": 0, "suppressed": True, "not_applicable": True}},
            "negative_coping": {"headline": {"value": 0.1, "n_valid": 200, "suppressed": False, "not_applicable": False}},
        }
    },
    "part_2": {
        "claims_funnel": {
            "filed_claim": {"n": 100, "pct_of_event_base": 0.2, "suppressed": False, "not_applicable": False},
            "experienced_event": {"n": 400, "n_valid": 400, "suppressed": False, "not_applicable": False},
            "claim_paid": {"value": 0.7, "n": 70, "n_valid": 100, "suppressed": False, "not_applicable": False},
            "payout_adequacy": {"n_valid": 80, "suppressed": False, "not_applicable": False},
        }
    },
    "part_1": {
        "metrics": {
            "coverage_understanding": {"headline": {"value": 0.5, "n_valid": 500, "suppressed": False, "not_applicable": False}},
            "claim_process_understanding": {"headline": {"value": 0.5, "n_valid": 500, "suppressed": False, "not_applicable": False}},
            # Deliberately not 0.5 like the metric above -- a tie here would
            # trip generation/executive_summary.py's
            # _disambiguate_tied_percentages() and this fixture's exec-summary
            # table test expects plain 1-decimal formatting.
            "worth_premium": {"headline": {"value": 0.62, "n_valid": 500, "suppressed": False, "not_applicable": False}},
            "renewal_intent": {"headline": {"value": 0.5, "n_valid": 500, "suppressed": False, "not_applicable": False}},
            "product_understanding": {"headline": {"value": None, "n_valid": 0, "suppressed": True, "not_applicable": True}},
        }
    },
}


class TestAddExecutiveSummary:
    def _doc_text(self, analysis: dict, qual: dict) -> str:
        from docx import Document
        doc = Document()
        _add_executive_summary(doc, analysis, qual)
        return "\n".join(p.text for p in doc.paragraphs)

    def test_headline_numbers_table_rendered(self):
        from docx import Document
        doc = Document()
        # report_scope="lacro" -- worth_premium's base_label is null for LACRO
        # (100% Health, no restriction), non-null for every other scope; see
        # generation/report_spec.yaml's executive_summary.metrics and R-002.
        _add_executive_summary(doc, {"parts": _EXEC_SUMMARY_PARTS, "meta": {"report_scope": "lacro"}}, {})
        table = doc.tables[0]
        rows = [[c.text for c in row.cells] for row in table.rows]
        assert rows[0] == ["Metric", "Value", "N", "Base"]
        assert ["First-Time Access to Insurance", "80.0%", "500", ""] in rows
        assert ["Worth the Premium", "62.0%", "500", ""] in rows
        assert ["Claim Process Understanding", "50.0%", "500", ""] in rows
        assert ["Children's Wellbeing Improved", "35.0%", "400", "clients with children in household"] in rows

    def test_caveat_box_lists_not_applicable_metric(self):
        body = self._doc_text({"parts": _EXEC_SUMMARY_PARTS}, {})
        assert "Data Availability" in body
        assert "Combined Product Understanding" in body

    def test_no_qual_data_omits_findings_actions_but_keeps_numbers(self):
        body = self._doc_text({"parts": _EXEC_SUMMARY_PARTS}, {})
        assert "Top Findings" not in body
        assert "Recommended Actions" not in body

    def test_qual_data_adds_narrative_findings_and_actions(self):
        qual = {
            "executive_summary": "Clients value fast payouts.",
            "top_findings": ["Finding A", "Finding B", "Finding C", "Finding D"],
            "top_actions": ["Action A", "Action B", "Action C"],
        }
        body = self._doc_text({"parts": _EXEC_SUMMARY_PARTS}, qual)
        assert "Clients value fast payouts." in body
        assert "Top Findings" in body
        assert "Finding A" in body and "Finding C" in body
        # Only the top 3 are shown even if the model returns more.
        assert "Finding D" not in body
        assert "Recommended Actions" in body
        assert "Action A" in body

    def test_recommended_actions_use_a_distinct_numbering_style(self):
        # C-018/R-013 (session-10): Top Findings and Recommended Actions
        # both used to render with style "List Number" -- one shared
        # numbering instance in Word, so Recommended Actions continued
        # Top Findings' count (1-3) as 4-6 instead of restarting at 1.
        # "List Number 2" is a distinct built-in style with its own numId
        # (confirmed: python-docx's default template gives ListNumber/
        # ListNumber2 numId 5/6), so it restarts independently.
        from docx import Document
        doc = Document()
        qual = {
            "executive_summary": "Summary.",
            "top_findings": ["Finding A", "Finding B"],
            "top_actions": ["Action A", "Action B"],
        }
        _add_executive_summary(doc, {"parts": _EXEC_SUMMARY_PARTS}, qual)
        finding_styles = {p.style.name for p in doc.paragraphs if p.text in ("Finding A", "Finding B")}
        action_styles = {p.style.name for p in doc.paragraphs if p.text in ("Action A", "Action B")}
        assert finding_styles == {"List Number"}
        assert action_styles == {"List Number 2"}
        assert finding_styles != action_styles


class TestAssembleExecutiveSummary:
    def _doc_text(self, doc_path) -> str:
        doc = Document(str(doc_path))
        return "\n".join(p.text for p in doc.paragraphs)

    def test_executive_summary_added_when_parts_present(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run_with_parts(
            tmp_path, "exec_run", {"country": "default", "n_total": 500}, _EXEC_SUMMARY_PARTS,
        )
        out = tmp_path / "out.docx"
        assemble([], {}, "exec_run", out)
        assert "Executive Summary" in self._doc_text(out)

    def test_executive_summary_precedes_about_this_survey(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        about = {"n_total": 500, "by_country": [], "product_mix": {"available": False},
                 "age": {}, "by_sex": [], "fieldwork": {"available": False}}
        parts = {**_EXEC_SUMMARY_PARTS, "about_survey": about}
        _make_run_with_parts(tmp_path, "order_run", {"country": "default", "n_total": 500}, parts)
        out = tmp_path / "out.docx"
        assemble([], {}, "order_run", out)
        body = self._doc_text(out)
        assert body.index("Executive Summary") < body.index("About This Survey")

    def test_qualitative_results_merged_in_when_present(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        qual = {
            "executive_summary": "A concise narrative.",
            "top_findings": ["F1", "F2", "F3"],
            "top_actions": ["A1", "A2", "A3"],
        }
        _make_run_with_parts(
            tmp_path, "qual_run", {"country": "default", "n_total": 500}, _EXEC_SUMMARY_PARTS, qual=qual,
        )
        out = tmp_path / "out.docx"
        assemble([], {}, "qual_run", out)
        body = self._doc_text(out)
        assert "A concise narrative." in body
        assert "F1" in body
        assert "A1" in body

    def test_no_parts_key_omits_executive_summary(self, tmp_path, monkeypatch):
        monkeypatch.setattr(assembler, "ROOT", tmp_path)
        _make_run(tmp_path, "no_parts_run", {"country": "default", "n_total": 100})
        out = tmp_path / "out.docx"
        assemble([], {}, "no_parts_run", out)
        assert "Executive Summary" not in self._doc_text(out)


# ---------------------------------------------------------------------------
# Part 10 -- Trend Comparison table (R-009: Sig. column replaced with
# Comparability; R-005: DR-exclusion scope note stated once near the table,
# not per row).
# ---------------------------------------------------------------------------

_TREND_PACKAGE = {
    "title": "Trend Comparison",
    "sections": {"insight": {"verbatims": []}},
    "visuals": [],
    "trend_scope_note": (
        "First-Time Access to Insurance and Client Satisfaction (NPS) report the five "
        "countries surveyed in both waves; Dominican Republic is new in 2026 and excluded "
        "from those row(s), with no 2025 counterpart."
    ),
    "scorecard": [
        {
            "label": "First-Time Access to Insurance",
            "group_a_label": "Current Wave", "group_a_value": "76.7%",
            "group_b_label": "Prior Wave", "group_b_value": "73.6%",
            "sig_p": None, "significant": False, "population": None,
            "sig_test_note": "Identical question wording and options in both waves.",
            "comparability": "clean",
        },
        {
            # session-5 (LM3, per Lorenz): both wave values shown, not
            # "NOT COMPARABLE" -- the instrument changed but a real figure
            # exists on both sides (48.9% in 2025, 44.5% in 2026).
            "label": "Access to Alternatives (Difficult)",
            "group_a_label": "Current Wave", "group_a_value": "44.5%",
            "group_b_label": "Prior Wave", "group_b_value": "48.9%",
            "sig_p": None, "significant": False, "population": None,
            "sig_test_note": "Indicative only, not a rigorous comparison: the 2025 instrument "
                              "offered 4 forced-choice options; 2026 adds a neutral midpoint. "
                              "Both figures are shown for reference, not as a tested change -- "
                              "do not use comparative language to describe the difference "
                              "between them.",
            "comparability": "indicative",
        },
        {
            # product_understanding's genuinely different case: 2026 truly
            # has no figure (not_applicable), but 2025's real value is
            # still shown, not suppressed to "NOT COMPARABLE" either.
            "label": "Product Understanding",
            "group_a_label": "Current Wave", "group_a_value": "NOT APPLICABLE",
            "group_b_label": "Prior Wave", "group_b_value": "20.0%",
            "sig_p": None, "significant": False, "population": None,
            "sig_test_note": "Not comparable to the prior wave: the 2025 instrument used one "
                              "combined 6-option question. Both figures are shown for reference, "
                              "not as a tested change -- do not use comparative language to "
                              "describe the difference between them. This wave's format has no "
                              "figure for this indicator at all, in the form it is defined from "
                              "-- only the prior wave's own figure is available.",
            "comparability": "not_comparable",
        },
    ],
}


class TestBuildPart10TrendTable:
    def _render(self, package=None):
        doc = Document()
        build_part_10(doc, package or _TREND_PACKAGE, {"narrative": "x", "insight": "y"})
        return doc

    def test_headers_are_chronological_indicator_years_comparability(self):
        # Chronological (2025 then 2026), not "current wave first" --
        # Lorenz's explicit call: left-to-right chronology is what a
        # reader expects from a trend table.
        doc = self._render()
        table = doc.tables[0]
        headers = [c.text for c in table.rows[0].cells]
        assert headers == ["Indicator", "2025", "2026", "Comparability"]

    def test_no_sig_column_or_significance_caption_anywhere(self):
        doc = self._render()
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                assert "Sig" not in cell.text
        body = "\n".join(p.text for p in doc.paragraphs)
        assert "z-test" not in body
        assert "p < 0.05" not in body

    def test_comparability_column_shows_status_word_only(self):
        doc = self._render()
        table = doc.tables[0]
        data_rows = [[c.text for c in row.cells] for row in table.rows[1:]]
        # Columns are now [label, 2025 (group_b), 2026 (group_a), comparability].
        assert ["First-Time Access to Insurance ‡", "73.6%", "76.7%", "Clean"] in data_rows
        assert ["Access to Alternatives (Difficult) ‡", "48.9%", "44.5%", "Indicative"] in data_rows
        assert ["Product Understanding ‡", "20.0%", "NOT APPLICABLE", "Not comparable"] in data_rows
        # The reason text itself must NOT be in the table cell -- it wraps
        # badly there and belongs in the footnote instead (Lorenz's request).
        for row in data_rows:
            assert "instrument" not in row[3]
            assert "wording" not in row[3]

    def test_indicative_and_not_comparable_rows_show_real_prior_values(self):
        # session-5 (LM3, per Lorenz): the whole point of the fix -- a
        # non-"clean" row must never render the literal string
        # "NOT COMPARABLE" when a real number exists for that wave.
        doc = self._render()
        table = doc.tables[0]
        for row in table.rows[1:]:
            cells = [c.text for c in row.cells]
            assert "NOT COMPARABLE" not in cells

    def test_footnotes_state_no_comparative_language_for_non_clean_rows(self):
        body = "\n".join(p.text for p in self._render().paragraphs)
        assert "do not use comparative language" in body

    def test_reason_appears_as_a_footnote_not_in_the_cell(self):
        body = "\n".join(p.text for p in self._render().paragraphs)
        assert "‡ First-Time Access to Insurance: Identical question wording" in body
        assert "‡ Access to Alternatives (Difficult): Indicative only" in body
        assert "‡ Product Understanding: Not comparable to the prior wave" in body
        assert "only the prior wave's own figure is available" in body

    def test_scope_note_rendered_once_near_the_table(self):
        body = "\n".join(p.text for p in self._render().paragraphs)
        assert body.count("Dominican Republic is new in 2026") == 1
        assert "five countries surveyed in both waves" in body

    def test_no_scope_note_omits_the_paragraph_entirely(self):
        package = dict(_TREND_PACKAGE, trend_scope_note="")
        body = "\n".join(p.text for p in self._render(package).paragraphs)
        assert "Dominican Republic" not in body
