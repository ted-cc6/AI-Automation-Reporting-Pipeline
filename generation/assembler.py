"""generation/assembler.py

Phase 4: Build the final .docx using python-docx. No Gemini calls here.
"""
import json
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from analysis_engine.country_config import DEFAULT_COUNTRY
from generation.executive_summary import data_availability_caveats, headline_numbers
from generation.writer import _lacro_scoped
from report_scopes import LACRO_SHORT_LABEL, REPORT_SCOPES
from utils import format_period_label, format_p_value

log = logging.getLogger(__name__)

ROOT = Path(__file__).parent.parent

# VisionFund brand orange (from the logo) — replaces python-docx's default
# blue heading theme colors (Title/Heading 1-4 come out of the box as shades
# of blue, e.g. #4F81BD, since nothing here set a color before).
VFI_ORANGE = RGBColor(0xF3, 0x66, 0x1F)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_default_font(doc, name: str, size: int):
    style = doc.styles["Normal"]
    font  = style.font
    font.name = name
    font.size = Pt(size)


def _apply_brand_heading_color(doc):
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            doc.styles[style_name].font.color.rgb = VFI_ORANGE
        except KeyError:
            pass


def _load_json(path: Path) -> "dict | None":
    """Best-effort JSON read. Returns None if the file is missing or
    unreadable rather than raising."""
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return None


def _load_analysis_meta(run_id: str) -> dict:
    """Best-effort read of runs/{run_id}/analysis_results.json's meta block, for
    the cover page. Returns {} if the file is missing or unreadable rather than
    raising -- the cover page falls back to a plain title in that case."""
    data = _load_json(ROOT / "runs" / run_id / "analysis_results.json")
    return (data or {}).get("meta", {})


def _load_full_analysis(run_id: str) -> dict:
    """Best-effort read of the full analysis_results.json -- used for
    sections that read across multiple parts (e.g. the executive summary's
    headline numbers, see generation/executive_summary.py) rather than one
    part_key. Returns {} if missing/unreadable."""
    return _load_json(ROOT / "runs" / run_id / "analysis_results.json") or {}


def _load_analysis_part(run_id: str, part_key: str) -> "dict | None":
    """Best-effort read of one parts.{part_key} block from analysis_results.json
    -- used for sections rendered directly from analysis data with no LLM call
    involved (e.g. About This Survey), so they don't need an orchestrator.py
    package or a report_spec.yaml entry at all. Returns None if the file, the
    part, or the JSON itself is missing/unreadable."""
    data = _load_json(ROOT / "runs" / run_id / "analysis_results.json")
    if data is None:
        return None
    return (data.get("parts") or {}).get(part_key)


def _load_qualitative_results(run_id: str) -> dict:
    """Best-effort read of runs/{run_id}/qualitative_results.json -- used for
    the executive summary's top_findings/top_actions/executive_summary prose
    (see qualitative/llm_call.py's Task 7). Returns {} if missing/unreadable
    (e.g. this run has no qualitative pass), so the executive summary falls
    back to headline numbers + caveat box only."""
    return _load_json(ROOT / "runs" / run_id / "qualitative_results.json") or {}


def _format_profile(profile: dict) -> str:
    parts = []
    if profile.get("sex"):
        parts.append(profile["sex"])
    if profile.get("age"):
        parts.append(f"age {profile['age']}")
    if profile.get("branch") and profile.get("country"):
        parts.append(f"{profile['branch']}, {profile['country']}")
    elif profile.get("branch"):
        parts.append(profile["branch"])
    elif profile.get("country"):
        parts.append(profile["country"])
    flags = []
    if profile.get("is_claimant"):
        flags.append("claimant")
    if profile.get("is_caregiver"):
        flags.append("caregiver")
    if flags:
        parts.append(", ".join(flags))
    return " | ".join(parts) if parts else "anonymous"


def _add_heading(doc, text: str, level: int):
    doc.add_heading(text, level=level)


def _add_paragraph(doc, text: str, style: str = None):
    if not text:
        return
    if style:
        try:
            doc.add_paragraph(text, style=style)
        except KeyError:
            doc.add_paragraph(text)
    else:
        doc.add_paragraph(text)


def _add_insight_box(doc, insight_text: str, verbatims: list):
    # R-010 (docs/report_spec.md, session-10): a section with no verbatim
    # candidates used to still render the "Key Qualitative Insights"
    # heading plus insight_text -- and insight_text, written by the same
    # LLM call that had nothing to quote, would itself narrate that
    # absence (e.g. "...have not yet been analyzed for this report"),
    # tripping C-014 (LM4/LM11) on top of C-015. Omitting the whole block
    # when no verbatim survives selection removes both failure modes at
    # once: there is nothing left to narrate the absence of.
    valid_verbatims = [v for v in verbatims[:3] if v.get("text", "")]
    if not valid_verbatims:
        return
    _add_heading(doc, "Key Qualitative Insights", level=3)
    if insight_text:
        _add_paragraph(doc, insight_text)
    for v in valid_verbatims:
        text    = v.get("text", "")
        profile = _format_profile(v.get("profile", {}))
        try:
            p = doc.add_paragraph(f'"{text}"', style="Quote")
        except KeyError:
            p = doc.add_paragraph(f'"{text}"')
            p.paragraph_format.left_indent = Inches(0.5)
        attr = doc.add_paragraph(f"({profile})")
        attr.paragraph_format.left_indent = Inches(0.5)
        if attr.runs:
            attr.runs[0].italic = True


def _group_header(groups: dict, key: str, fallback_label: str) -> str:
    """Column header for a scorecard table group, e.g. 'Claimant (n=1,922)',
    or 'Claimant (filed, n=55)' when the group dict carries an optional
    "qualifier" (R-011, docs/report_spec.md, session-10) -- Part 6 uses this
    to state each group's population inline in the header itself, rather
    than a single word ("Non-Filer") a reader has to infer meaning from.
    Parts 5/7 don't set "qualifier", so this is unchanged for them."""
    g = groups.get(key, {})
    label = g.get("label", fallback_label)
    n = g.get("n")
    qualifier = g.get("qualifier")
    if isinstance(n, int):
        if qualifier:
            return f"{label} ({qualifier}, n={n:,})"
        return f"{label} (n={n:,})"
    return label


def _add_drivers_table(doc, section: dict):
    """Renders a Spearman correlation table (Part 4's factors associated with
    satisfaction, Part 5's factors associated with child wellbeing) -- shared
    since both sections use the identical shape."""
    drivers_data  = section.get("drivers_data", [])
    drivers_table = section.get("drivers_table", {})
    if not drivers_data:
        return
    headers = drivers_table.get("headers", ["Factor", "ρ (Spearman)", "p-value", "N"])
    rows = []
    for d in drivers_data:
        if d.get("not_applicable"):
            rows.append([d["label"], "NOT APPLICABLE", "NOT APPLICABLE", "NOT APPLICABLE"])
        elif d["suppressed"]:
            rows.append([d["label"], "SUPPRESSED", "SUPPRESSED", "SUPPRESSED"])
        else:
            rho_str = f"{d['rho']:+.3f}" if d["rho"] is not None else "?"
            p_str   = format_p_value(d["p_value"])
            n_str   = str(d["n_valid"]) if d["n_valid"] is not None else "?"
            rows.append([d["label"], rho_str, p_str, n_str])
    _add_table(doc, headers, rows)
    # R-035 (docs/report_spec.md, session-11): a reader looking at the table
    # itself, not just the methodology appendix, is where a causal reading is
    # most tempting -- this note is deterministic, code-rendered, not left to
    # the writer prompt to remember to say every time.
    association_note = drivers_table.get("association_note")
    if association_note:
        _add_paragraph(doc, association_note)
    # Full ρ/p-value/N methodology explanation lives once in the Methodology
    # Notes appendix (see _add_methodology_appendix()) rather than repeated
    # verbatim under both Part 4's and Part 5's drivers tables.
    _add_paragraph(doc, "See \"Appendix: Methodology Notes\" for how to read ρ (rho), p-value, and N in this table.")


def _add_image_or_placeholder(doc, visual_info: dict):
    if visual_info.get("exists") and visual_info.get("path"):
        try:
            doc.add_picture(str(visual_info["path"]), width=Inches(5.5))
            p = doc.add_paragraph(visual_info.get("caption", ""), style="Caption")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            p = doc.add_paragraph(
                f'[VISUAL: could not insert {visual_info["file"]} ({visual_info.get("caption", "")})]'
            )
            if p.runs:
                p.runs[0].italic = True
    else:
        p = doc.add_paragraph(
            f'[VISUAL PENDING: {visual_info["file"]} ({visual_info.get("caption", "")})]'
        )
        if p.runs:
            p.runs[0].italic = True


_PROTECTION_FLAG_LABELS = {
    "mis_selling": "Mis-selling",
    "premium_without_consent": "Premium deducted without consent",
    "coercion": "Coercion to purchase",
    "false_information": "False information",
    "unfair_claim_denial": "Unfair claim denial",
    "staff_misconduct": "Staff misconduct",
    "data_privacy": "Data privacy concern",
}

_SEVERITY_ORDER = ["high", "medium", "low"]


def _add_bold_paragraph(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def _protection_flag_ref(flag: dict) -> str:
    """Reader-friendly, actually-traceable reference for a protection-flag
    case: client_id plus branch, matching how the client protection team
    looks up a client in their own systems -- a row_id/row-index number
    means nothing to them and doesn't survive re-runs of the pipeline
    (row order can shift; client_id doesn't)."""
    profile = flag.get("profile") or {}
    client_id = profile.get("client_id")
    branch = profile.get("branch")
    if client_id and branch:
        return f"{client_id}, {branch}"
    if client_id:
        return client_id
    # Fall back to the internal id only when no profile could be resolved
    # (e.g. a row_id that no longer maps into the survey dataframe).
    return flag.get("id", "") or "unknown"


def _compute_severity_counts(protection_flags: list) -> dict:
    """Single source of truth for severity tallies (R-017). Both the
    in-body summary and the appendix call this on the same canonical
    protection_flags list rather than each maintaining its own counting
    loop -- identical input to a pure function guarantees identical
    output, so the two renderings can't drift apart."""
    by_severity: dict[str, int] = {}
    for flag in protection_flags:
        sev = (flag.get("severity") or "unspecified").lower()
        by_severity[sev] = by_severity.get(sev, 0) + 1
    return by_severity


def _add_protection_signals_summary(doc, protection_flags: list, severity_counts: dict) -> None:
    """Short in-body summary (counts by severity) inline in Part 2 -- the
    full itemized list with per-case reason text and client references for
    follow-up moves to "Appendix: Client Protection Signals" instead (see
    _add_protection_signals_annex()), so a run with many flags doesn't let a
    long itemized list dominate Part 2's narrative flow."""
    if not protection_flags:
        return
    _add_heading(doc, "Client Protection Signals", level=4)

    ordered_keys = [s for s in _SEVERITY_ORDER if s in severity_counts]
    ordered_keys += [s for s in severity_counts if s not in _SEVERITY_ORDER]
    counts_text = ", ".join(f"{severity_counts[s]} {s}" for s in ordered_keys)

    n_total = len(protection_flags)
    _add_paragraph(
        doc,
        f"{n_total} client-reported protection concern{'s' if n_total != 1 else ''} "
        f"{'were' if n_total != 1 else 'was'} identified during this survey ({counts_text} "
        "severity). See \"Appendix: Client Protection Signals\" for the full itemized list "
        "with client references for follow-up by the client protection team."
    )


def _add_protection_signals_annex(doc, protection_flags: list, severity_counts: dict) -> None:
    """Full itemized client-protection flag list (reason text + client
    reference per case), grouped by severity -- rendered once as a document
    appendix instead of inline in Part 2 (see _add_protection_signals_summary()
    for the short in-body version Part 2 shows instead)."""
    if not protection_flags:
        return
    doc.add_heading("Appendix: Client Protection Signals", level=1)
    _add_paragraph(
        doc,
        "The following client-reported concerns were identified for follow-up "
        "by the client protection team."
    )

    by_severity: dict[str, list] = {}
    for flag in protection_flags:
        sev = (flag.get("severity") or "unspecified").lower()
        by_severity.setdefault(sev, []).append(flag)
    ordered_keys = [s for s in _SEVERITY_ORDER if s in severity_counts]
    ordered_keys += [s for s in severity_counts if s not in _SEVERITY_ORDER]

    for sev in ordered_keys:
        _add_bold_paragraph(doc, f"{sev.capitalize()} severity")
        for flag in by_severity.get(sev, []):
            flag_type = flag.get("flag_type", "")
            label = _PROTECTION_FLAG_LABELS.get(flag_type, flag_type.replace("_", " ").capitalize())
            reason = (flag.get("reason") or "").strip()
            ref = _protection_flag_ref(flag)
            note = "; same client, multiple concerns" if flag.get("same_client_multiple_concerns") else ""
            _add_paragraph(doc, f"{label}: {reason} ({ref}{note})", style="List Bullet")


_SPEARMAN_METHODOLOGY_INTRO = (
    "ρ (rho) is the Spearman rank correlation coefficient, a value from -1 to +1 showing "
    "how strongly two things move together: values near +1 or -1 indicate a strong "
    "association, values near 0 indicate little to none. Spearman correlation is "
    "calculated on the rank order of responses rather than their raw values, which makes "
    "it well suited to Likert-scale survey questions like the ones in Parts 4 and 5. p "
    "(the p-value) shows how likely an association this strong could appear by chance alone "
    "if no real association existed; p<0.05 is the conventional threshold for statistical "
    "significance. N is the number of respondents included in each specific calculation, "
    "which can vary by factor because not every question was asked of every client{n_example}. "
    "These are observed associations in cross-sectional survey data, collected at a single "
    "point in time; they do not establish that one factor is responsible for another."
)

# The Renewal Intent/Vietnam example is only true, and only a driver at all,
# for a report scope where renewal_intent actually appears in the drivers
# table (Africa+Vietnam) -- LACRO never asks that question at all (see
# report_scopes.py), so its drivers table never has a renewal_intent row to
# illustrate the point with, and the example was flatly false for a LACRO
# report ("in the Africa/Vietnam portfolio, Renewal Intent...") that has no
# Africa or Vietnam clients in it at all.
_SPEARMAN_N_EXAMPLE_DEFAULT = (
    " (for example, in the Africa/Vietnam portfolio, Renewal Intent was asked only of "
    "Vietnam's crop-insurance clients)"
)

_SPEARMAN_METHODOLOGY_SCORING_NOTE_DEFAULT = (
    "Financial Stress, Coverage Understanding, Claim Process Understanding, Worth Premium, "
    "Renewal Intent, and Confidence in Payout are all scored so that a LOWER number is the "
    "more positive response (e.g. 1 = \"Definitely would renew\"). A negative correlation "
    "with the outcome variable (client satisfaction, NPS, in Part 4; child wellbeing in "
    "Part 5) for any of these factors therefore corresponds to a POSITIVE association in "
    "plain terms (e.g. stronger renewal intent occurring alongside a better outcome), not a "
    "negative one. NPS itself runs the opposite way (0=worst, 10=best): when NPS appears as "
    "the outcome (Part 4), a negative correlation with it is what corresponds to the positive "
    "association described above; when NPS appears as a FACTOR associated with child wellbeing "
    "(Part 5 only), a positive correlation there likewise corresponds to a positive association."
)

# LACRO never asks Renewal Intent (see _SPEARMAN_N_EXAMPLE_DEFAULT's comment
# above) -- its drivers table has no such row, so naming it as a scoring
# example would be exactly as false here as the sibling N-variability
# example this file already guards against for the same reason. Uses a
# scope-neutral phrasing for the scoring-direction example instead of
# naming a specific driver that does not exist in a LACRO report's data.
_SPEARMAN_METHODOLOGY_SCORING_NOTE_LACRO = (
    "Financial Stress, Coverage Understanding, Claim Process Understanding, Worth Premium, "
    "and Confidence in Payout are all scored so that a LOWER number is the more positive "
    "response (e.g. 1 = the most positive answer option for that question). A negative "
    "correlation with the outcome variable (client satisfaction, NPS, in Part 4; child "
    "wellbeing in Part 5) for any of these factors therefore corresponds to a POSITIVE "
    "association in plain terms (e.g. a more positive response occurring alongside a better "
    "outcome), not a negative one. NPS itself runs the opposite way (0=worst, 10=best): when "
    "NPS appears as the outcome (Part 4), a negative correlation with it is what corresponds "
    "to the positive association described above; when NPS appears as a FACTOR associated "
    "with child wellbeing (Part 5 only), a positive correlation there likewise corresponds to "
    "a positive association."
)


def _spearman_methodology_note(meta: dict) -> str:
    # Was gated on report_scope == "lacro" alone, which misses the legacy
    # dataset_schema == "larco" upload path (the frontend never offers the
    # report_scope picker for that schema -- CupboardWeekApp.tsx -- so a real
    # run through it has report_scope == None). Use the same is-this-a-LARCO-
    # report check every other LACRO-conditional branch in this codebase uses
    # (_lacro_scoped(), writer.py; is_lacro_report, run_analysis.py) so this
    # is the one place that can't drift back out of sync with the others.
    lacro = _lacro_scoped(meta)
    n_example = "" if lacro else _SPEARMAN_N_EXAMPLE_DEFAULT
    intro = _SPEARMAN_METHODOLOGY_INTRO.format(n_example=n_example)
    scoring_note = (
        _SPEARMAN_METHODOLOGY_SCORING_NOTE_LACRO if lacro
        else _SPEARMAN_METHODOLOGY_SCORING_NOTE_DEFAULT
    )
    return f"{intro}\n\n{scoring_note}"


def _add_methodology_appendix(doc, meta: "dict | None" = None):
    """Single shared explanation of the Spearman ρ/p-value/N methodology used
    by both Part 4's satisfaction drivers table and Part 5's child wellbeing
    drivers table -- previously duplicated verbatim (~200 words) under each
    table; both now print a one-line pointer to this appendix instead (see
    _add_drivers_table())."""
    doc.add_heading("Appendix: Methodology Notes", level=1)
    _add_heading(doc, "Spearman Rank Correlation (Parts 4 & 5 Factors Tables)", level=2)
    for para in _spearman_methodology_note(meta or {}).split("\n\n"):
        _add_paragraph(doc, para)


def _add_executive_summary(doc, analysis: dict, qual: dict):
    """Renders the Executive Summary: headline numbers (deterministic, from
    generation/executive_summary.py) + top findings/actions and a short
    narrative (from the qualitative pipeline's Task 7, qualitative_results.json
    -- see qualitative/llm_call.py) + a consolidated data-availability caveat
    box. No LLM call happens here directly -- the narrative/findings/actions
    were already generated by the qualitative pass; this only merges and
    renders what already exists, so a run with no qualitative_results.json
    still gets a useful headline-numbers-only executive summary rather than
    nothing at all."""
    _add_heading(doc, "Executive Summary", level=1)

    rows = headline_numbers(analysis)
    if rows:
        _add_table(
            doc, ["Metric", "Value", "N", "Base"],
            [[r["label"], r["value"], r["n"] if r["n"] is not None else "", r["base_label"]] for r in rows],
        )

    exec_prose = (qual or {}).get("executive_summary", "")
    if exec_prose:
        _add_paragraph(doc, exec_prose)

    top_findings = (qual or {}).get("top_findings", [])[:3]
    if top_findings:
        _add_heading(doc, "Top Findings", level=3)
        for finding in top_findings:
            _add_paragraph(doc, finding, style="List Number")

    top_actions = (qual or {}).get("top_actions", [])[:3]
    if top_actions:
        _add_heading(doc, "Recommended Actions", level=3)
        # C-018/R-013: "List Number" is one shared numbering instance for the
        # whole document -- reusing it here continued Top Findings' count
        # (1-3) into Recommended Actions (4-6) instead of restarting at 1.
        # "List Number 2" is a distinct built-in Word style with its own
        # independent numId (confirmed: python-docx's default template
        # gives ListNumber/ListNumber2 numId 5/6 respectively), so it
        # restarts cleanly without any raw-XML numbering surgery.
        for action in top_actions:
            _add_paragraph(doc, action, style="List Number 2")

    caveats = data_availability_caveats(analysis)
    if caveats:
        _add_heading(doc, "Data Availability", level=3)
        _add_paragraph(
            doc,
            "This report's survey instrument does not collect the following, so they do "
            "not appear elsewhere in this report: " + "; ".join(caveats) + "."
        )


def _add_data_notes_subsection(doc, data_notes: "dict | None", quality_flags: "list | None" = None):
    """Short, deterministic "Data Notes" subsection stating exclusion counts,
    the dedup rule, and any active data-quality flags -- surfaced inside the
    report itself, not only in the side screening_report.md diagnostic file
    nobody reading the report ever sees. Source: data_loader_screening.py's
    build_screening_summary() (the "data_notes" key) and data_quality_flags.
    get_flags() (the "data_quality_flags" key), both top-level in
    analysis_results.json. None of the dedup/exclusion counts were ever a
    correctness bug (exact-content duplicates are genuinely dropped,
    client-ID reuse is genuinely kept and flagged) -- the report stating
    nothing about it is what made it look like one.
    """
    if not data_notes and not quality_flags:
        return
    _add_heading(doc, "Data Notes", level=3)

    data_notes = data_notes or {}
    removed = data_notes.get("removed", {})
    n_start = data_notes.get("n_start")
    n_end = data_notes.get("n_end")
    lines = []
    if n_start is not None and n_end is not None:
        lines.append(f"{n_start:,} client responses were exported; {n_end:,} are covered by this report.")
    if removed.get("duplicate_submission"):
        lines.append(
            f"{removed['duplicate_submission']} duplicate-submission row(s) removed "
            f"({data_notes.get('rules', {}).get('duplicate_submission', '')})"
        )
    if removed.get("non_consenting"):
        lines.append(f"{removed['non_consenting']} row(s) removed for declining consent.")
    if removed.get("test_qa"):
        lines.append(f"{removed['test_qa']} test/QA row(s) removed.")
    if removed.get("unselected_region"):
        lines.append(
            f"{removed['unselected_region']:,} row(s) outside this report's scope "
            f"({data_notes.get('report_scope')!r}) removed."
        )
    if removed.get("unselected_country"):
        lines.append(
            f"{removed['unselected_country']:,} row(s) outside the selected country "
            f"({data_notes.get('target_country')!r}) removed."
        )
    for line in lines:
        _add_paragraph(doc, line)

    n_warn = data_notes.get("client_id_reuse_warnings", 0)
    if n_warn:
        _add_paragraph(
            doc,
            f"{n_warn} client ID(s) appear more than once with different answers "
            f"({data_notes.get('rules', {}).get('client_id_reuse_not_dropped', '')}) "
            "Both records are kept as-is pending field-team reconciliation."
        )

    uuid_pairs = data_notes.get("uuid_duplicate_pairs") or []
    if uuid_pairs:
        n_high = sum(1 for p in uuid_pairs if p.get("severity") == "high")
        _add_paragraph(
            doc,
            f"{len(uuid_pairs)} pair(s) of records share a device-assigned ID but are not "
            f"identical ({n_high} at high similarity, suggesting partial answer carryover "
            "between two interviews on the same device). Both records in each pair are kept "
            "as-is pending field-team reconciliation."
        )

    for flag in (quality_flags or []):
        _add_paragraph(doc, f"Data quality flag ({flag.get('id')}): {flag.get('note', '')}")


def _add_about_survey_section(doc, about: dict, data_notes: "dict | None" = None,
                               quality_flags: "list | None" = None):
    """Renders analysis_engine/sections/about_survey.py's output directly --
    no LLM call, no report_spec.yaml entry, no orchestrator.py package. Pure
    descriptive statistics (country/product/age/sex/fieldwork dates), so
    there is nothing here for a model to get wrong or need word-limited."""
    _add_heading(doc, "About This Survey", level=1)

    n_total = about.get("n_total", 0)
    fieldwork = about.get("fieldwork", {})
    if fieldwork.get("available"):
        _add_paragraph(
            doc,
            f"This report is based on {n_total:,} client responses collected between "
            f"{fieldwork['start_date']} and {fieldwork['end_date']}."
        )
    else:
        _add_paragraph(doc, f"This report is based on {n_total:,} client responses.")

    by_country = about.get("by_country", [])
    if by_country:
        _add_heading(doc, "Respondents by Country", level=3)
        _add_table(
            doc, ["Country", "N", "%"],
            [[row["country"], row["n"], f"{row['pct'] * 100:.1f}%"] for row in by_country],
        )

    product_mix = about.get("product_mix", {})
    if product_mix.get("available"):
        _add_heading(doc, "Respondents by Product", level=3)
        _add_table(
            doc, ["Product", "N", "%"],
            [[row["product"].replace("_", " ").title(), row["n"], f"{row['pct'] * 100:.1f}%"]
             for row in product_mix.get("distribution", [])],
        )
    elif product_mix.get("reason"):
        log.info(f"About This Survey: product mix omitted — {product_mix['reason']}")

    age = about.get("age", {})
    if age.get("n_valid"):
        _add_heading(doc, "Age", level=3)
        _add_paragraph(
            doc,
            f"Mean age {age['mean']:.1f} years (range {age['min']} to {age['max']}), "
            f"based on {age['n_valid']:,} respondents with a recorded age."
        )

    by_sex = about.get("by_sex", [])
    if by_sex:
        _add_heading(doc, "Respondents by Sex", level=3)
        _add_table(
            doc, ["Sex", "N", "%"],
            [[row["sex"], row["n"], f"{row['pct'] * 100:.1f}%"] for row in by_sex],
        )

    _add_data_notes_subsection(doc, data_notes, quality_flags)


def _add_generation_failure_notice(doc, error: str):
    p = doc.add_paragraph(
        f"[NARRATIVE GENERATION FAILED: Gemini call did not succeed after retries "
        f"({error}). This section requires manual write-up before publishing.]"
    )
    if p.runs:
        p.runs[0].italic = True
        p.runs[0].bold = True


def _add_table(doc, headers: list, rows: list) -> object:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val) if val is not None else ""
    return table


# ---------------------------------------------------------------------------
# Part 1
# ---------------------------------------------------------------------------

def build_part_1(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 1: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    for s_key, vis_idx in [("s1_1", 0), ("s1_2", 1)]:
        s_data = sections.get(s_key, {})
        _add_heading(doc, s_data.get("label", s_key), level=2)
        _add_paragraph(doc, texts.get(s_key, ""))
        if vis_idx < len(visuals):
            _add_image_or_placeholder(doc, visuals[vis_idx])

    s2b = sections.get("s1_2b", {})
    if s2b:
        _add_heading(doc, s2b.get("label", "Preferred Claims Channel"), level=2)
        _add_paragraph(doc, texts.get("s1_2b", ""))

    s3 = sections.get("s1_3", {})
    _add_heading(doc, s3.get("label", "s1_3"), level=2)
    _add_paragraph(doc, texts.get("s1_3", ""))

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))


# ---------------------------------------------------------------------------
# Part 2
# ---------------------------------------------------------------------------

def build_part_2(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 2: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    # 2.1 — Claims Funnel
    s2_1 = sections.get("s2_1", {})
    _add_heading(doc, s2_1.get("label", "Claims Funnel"), level=2)
    _add_paragraph(doc, texts.get("s2_1", ""))

    ft = s2_1.get("funnel_table", {})
    if ft:
        metrics = s2_1.get("metrics", {})
        headers = ft.get("headers", ["Stage", "N", "Rate"])
        rows    = []
        for row_spec in ft.get("rows", []):
            n_key = row_spec.get("n_key", "")
            if n_key not in metrics:
                # orchestrator.py's extract_metrics() omits not_applicable metrics
                # entirely (e.g. LARCO has no "experienced insured event" gate, no
                # claim-paid outcome data) -- drop the row rather than a stage
                # label with blank N/Rate cells.
                continue
            n_val    = metrics.get(n_key, "")
            rate_val = metrics.get(row_spec.get("rate_key", ""), "")
            rows.append([row_spec.get("label", ""), n_val, rate_val])
        if rows:
            _add_table(doc, headers, rows)

    if len(visuals) > 0:
        _add_image_or_placeholder(doc, visuals[0])

    # 2.2 — No-claim reasons
    s2_2 = sections.get("s2_2", {})
    _add_heading(doc, s2_2.get("label", "Reasons for Not Claiming"), level=2)
    _add_paragraph(doc, texts.get("s2_2", ""))
    if len(visuals) > 1:
        _add_image_or_placeholder(doc, visuals[1])

    # 2.3 — Claim challenges
    s2_3 = sections.get("s2_3", {})
    _add_heading(doc, s2_3.get("label", "Claim Challenges"), level=2)
    _add_paragraph(doc, texts.get("s2_3", ""))
    if len(visuals) > 2:
        _add_image_or_placeholder(doc, visuals[2])

    prot_flags = s2_3.get("qualitative", {}).get("protection_flags", [])
    _add_protection_signals_summary(doc, prot_flags, _compute_severity_counts(prot_flags))

    # 2.4 — Payout outcomes
    s2_4 = sections.get("s2_4", {})
    _add_heading(doc, s2_4.get("label", "Payout Outcomes"), level=2)
    _add_paragraph(doc, texts.get("s2_4", ""))

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))


# ---------------------------------------------------------------------------
# Part 3
# ---------------------------------------------------------------------------

def build_part_3(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 3: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    s3_0 = sections.get("s3_0", {})
    if s3_0:
        _add_heading(doc, s3_0.get("label", "Reaching Clients with No Prior Insurance Access"), level=2)
        _add_paragraph(doc, texts.get("s3_0", ""))

    s3_1 = sections.get("s3_1", {})
    _add_heading(doc, s3_1.get("label", "Financial Stress and Coping"), level=2)
    _add_paragraph(doc, texts.get("s3_1", ""))
    if visuals:
        _add_image_or_placeholder(doc, visuals[0])

    s3_2 = sections.get("s3_2", {})
    _add_heading(doc, s3_2.get("label", "Confidence in Payout"), level=2)
    _add_paragraph(doc, texts.get("s3_2", ""))

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))


# ---------------------------------------------------------------------------
# Part 4
# ---------------------------------------------------------------------------

def build_part_4(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 4: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    s4_1 = sections.get("s4_1", {})
    _add_heading(doc, s4_1.get("label", "Net Promoter Score"), level=2)
    _add_paragraph(doc, texts.get("s4_1", ""))
    if visuals:
        _add_image_or_placeholder(doc, visuals[0])

    s4_2 = sections.get("s4_2", {})
    _add_heading(doc, s4_2.get("label", "Promoter and Detractor Themes"), level=2)
    _add_paragraph(doc, texts.get("s4_2", ""))
    if len(visuals) > 1:
        _add_image_or_placeholder(doc, visuals[1])

    s4_3 = sections.get("s4_3", {})
    if s4_3.get("drivers_data"):
        _add_heading(doc, s4_3.get("label", "Factors Associated with Satisfaction"), level=2)
        _add_paragraph(doc, texts.get("s4_3", ""))
        _add_drivers_table(doc, s4_3)

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))


# ---------------------------------------------------------------------------
# Part 5
# ---------------------------------------------------------------------------

def build_part_5(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 5: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    s5_1 = sections.get("s5_1", {})
    _add_heading(doc, s5_1.get("label", "Factors Associated with Child Wellbeing"), level=2)
    _add_paragraph(doc, texts.get("s5_1", ""))
    _add_drivers_table(doc, s5_1)

    if visuals:
        _add_image_or_placeholder(doc, visuals[0])

    s5_2 = sections.get("s5_2", {})
    _add_heading(doc, s5_2.get("label", "Healthcare Access and Medical Cost"), level=2)
    _add_paragraph(doc, texts.get("s5_2", ""))
    if len(visuals) > 1:
        _add_image_or_placeholder(doc, visuals[1])

    scorecard = package.get("scorecard", [])
    if scorecard:
        s5_3   = sections.get("s5_3", {})
        groups = package.get("groups", {})
        _add_heading(doc, s5_3.get("label", "Caregivers vs Non-Caregivers"), level=2)
        headers = ["Metric",
                   _group_header(groups, "caregiver", "Caregiver"),
                   _group_header(groups, "non_caregiver", "Non-Caregiver"),
                   "Sig.*"]
        rows, footnotes = [], []
        for row in scorecard:
            sig_mark = "*" if row["significant"] else ""
            label = row["label"]
            if row.get("population"):
                label += " †"
                footnotes.append(f"† {row['label']}: {row['population']}")
            if row.get("sig_test_note"):
                label += " ‡"
                footnotes.append(f"‡ {row['label']}: {row['sig_test_note']}")
            rows.append([label, row["group_a_value"], row["group_b_value"], sig_mark])
        _add_table(doc, headers, rows)
        _add_paragraph(doc, "* p < 0.05 (chi-squared or Fisher's exact test)")
        for fn in footnotes:
            _add_paragraph(doc, fn)
        _add_paragraph(doc, texts.get("s5_3", ""))

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))


# ---------------------------------------------------------------------------
# Part 6 — Claimant vs Did Not File
# ---------------------------------------------------------------------------

def build_part_6(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 6: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    scorecard = package.get("scorecard", [])
    if scorecard:
        groups = package.get("groups", {})
        headers = ["Metric",
                   _group_header(groups, "claimant", "Claimant (filed)"),
                   _group_header(groups, "non_claimant", "Did not file"),
                   "Sig.*"]
        rows, footnotes = [], []
        for row in scorecard:
            sig_mark = "*" if row["significant"] else ""
            label = row["label"]
            if row.get("population"):
                label += " †"
                footnotes.append(f"† {row['label']}: {row['population']}")
            if row.get("sig_test_note"):
                label += " ‡"
                footnotes.append(f"‡ {row['label']}: {row['sig_test_note']}")
            rows.append([label, row["group_a_value"], row["group_b_value"], sig_mark])
        _add_table(doc, headers, rows)
        # R-011 (docs/report_spec.md, session-10): "Non-Claimant" and
        # "Non-Filer" were each tried and found to let a reader infer the
        # wrong population on their own (see part_6.py's module docstring).
        # Stating the population directly, rather than trusting the column
        # header alone to carry it, is the actual fix.
        _add_paragraph(
            doc,
            "Both groups above are restricted to clients who reported an "
            "insured event in the past 12 months; neither describes "
            "clients who never faced a claimable event during that period."
        )
        _add_paragraph(doc, "* p < 0.05 (chi-squared or Fisher's exact test)")
        for fn in footnotes:
            _add_paragraph(doc, fn)

    if visuals:
        _add_image_or_placeholder(doc, visuals[0])

    _add_heading(doc, "Findings", level=2)
    _add_paragraph(doc, texts.get("narrative", ""))

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))


# ---------------------------------------------------------------------------
# Part 7 — Gender
# ---------------------------------------------------------------------------

def build_part_7(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 7: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    scorecard = package.get("scorecard", [])
    if scorecard:
        groups = package.get("groups", {})
        headers = ["Metric",
                   _group_header(groups, "female", "Female"),
                   _group_header(groups, "male", "Male"),
                   "Sig.*"]
        rows, footnotes = [], []
        for row in scorecard:
            sig_mark = "*" if row["significant"] else ""
            label = row["label"]
            if row.get("population"):
                label += " †"
                footnotes.append(f"† {row['label']}: {row['population']}")
            if row.get("sig_test_note"):
                label += " ‡"
                footnotes.append(f"‡ {row['label']}: {row['sig_test_note']}")
            rows.append([label, row["group_a_value"], row["group_b_value"], sig_mark])
        _add_table(doc, headers, rows)
        _add_paragraph(doc, "* p < 0.05 (chi-squared or Fisher's exact test)")
        for fn in footnotes:
            _add_paragraph(doc, fn)

    if visuals:
        _add_image_or_placeholder(doc, visuals[0])

    _add_heading(doc, "Findings", level=2)
    _add_paragraph(doc, texts.get("narrative", ""))

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))

    # Part 8 (Kling Index) is deliberately dashboard-only and never renders
    # here (see analysis_engine/sections/part_8.py's module docstring) --
    # without this line, a reader sees the part numbering jump straight from
    # 7 to whatever comes next (9, or 10 when a trend comparison moved to the
    # front) with nothing in the document itself explaining why 8 is
    # missing. This note is unconditional (Part 7 renders in every Cupboard
    # Week report), so the explanation appears regardless of which later
    # parts this particular run includes.
    _add_paragraph(
        doc,
        "Part 8 (Kling Index) is available on the analytics dashboard only and is not "
        "included in this document."
    )


# ---------------------------------------------------------------------------
# Part 9 — Additional Services (LARCO only)
# ---------------------------------------------------------------------------

def build_part_9(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 9: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    s9_1 = sections.get("s9_1", {})
    _add_heading(doc, s9_1.get("label", "Services Used"), level=2)
    _add_paragraph(doc, texts.get("s9_1", ""))
    dist = s9_1.get("distributions", {}).get("main", [])
    if dist:
        rows = [[d.get("option", d.get("value", "")), d.get("n", ""),
                 f"{d.get('pct', 0) * 100:.1f}%" if d.get("pct") is not None else ""]
                for d in dist]
        _add_table(doc, ["Service", "N", "%"], rows)
    if visuals:
        _add_image_or_placeholder(doc, visuals[0])

    s9_2 = sections.get("s9_2", {})
    _add_heading(doc, s9_2.get("label", "Service Helpfulness"), level=2)
    _add_paragraph(doc, texts.get("s9_2", ""))

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))


# ---------------------------------------------------------------------------
# Part 10 — Trend Comparison (renders whenever analysis_results.json has a
# parts.part_10 block -- see run_analysis.py's build_sections() and
# generation/orchestrator.py's orchestrate())
# ---------------------------------------------------------------------------

# R-009: status word only in the Comparability column (a full reason
# string wraps badly in a table cell); the reason itself renders as a
# per-row footnote instead -- see build_part_10()'s footnotes list.
_COMPARABILITY_DISPLAY = {
    "clean": "Clean",
    "indicative": "Indicative",
    "not_comparable": "Not comparable",
}


def build_part_10(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 10: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    scorecard = package.get("scorecard", [])
    if scorecard:
        # No groups() dict here (see orchestrator.py's _build_trend_data()
        # note) -- plain column headers, not _group_header()'s "(n=...)"
        # suffix, since a single per-table N would misrepresent rows that
        # each have their own base/suppression.
        #
        # R-009: "2025"/"2026" are hardcoded literals here, not derived
        # from a config value the way R-012 wants (prior_wave.year/
        # current_wave.year) -- R-012 is not in scope this session. A
        # future wave reusing this needs either R-012 implemented or these
        # two literals updated by hand.
        #
        # Column order is chronological (2025 then 2026), not "current
        # wave first" -- a reader expects left-to-right chronology in a
        # trend table; descending order invites misreading the direction
        # of change. row["group_a_value"]/["group_b_value"] stay
        # current/prior internally (unchanged from Parts 6/7's own
        # convention, and what writer.py's prompt text still calls
        # "Current Wave"/"Prior Wave" -- R-012, not this session), so the
        # two are swapped only here, at render time, not in the data.
        headers = ["Indicator", "2025", "2026", "Comparability"]
        rows, footnotes = [], []
        for row in scorecard:
            label = row["label"]
            comparability = _COMPARABILITY_DISPLAY.get(row.get("comparability"), "")
            if row.get("sig_test_note"):
                label += " ‡"
                footnotes.append(f"‡ {row['label']}: {row['sig_test_note']}")
            rows.append([label, row["group_b_value"], row["group_a_value"], comparability])
        scope_note = package.get("trend_scope_note", "")
        if scope_note:
            _add_paragraph(doc, scope_note)
        _add_table(doc, headers, rows)
        for fn in footnotes:
            _add_paragraph(doc, fn)

    if visuals:
        _add_image_or_placeholder(doc, visuals[0])

    _add_heading(doc, "Findings", level=2)
    _add_paragraph(doc, texts.get("narrative", ""))

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))


# ---------------------------------------------------------------------------
# Part 11 — Credit Life Module (renders whenever analysis_results.json has a
# parts.part_11 block -- see run_analysis.py's build_sections(), gated to
# report_scope=="africa")
# ---------------------------------------------------------------------------

def build_part_11(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 11: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    s11_1 = sections.get("s11_1", {})
    _add_heading(doc, s11_1.get("label", "Other Benefits Used"), level=2)
    _add_paragraph(doc, texts.get("s11_1", ""))
    dist = s11_1.get("distributions", {}).get("main", [])
    if dist:
        rows = [[d.get("option", d.get("value", "")), d.get("n", ""),
                 f"{d.get('pct', 0) * 100:.1f}%" if d.get("pct") is not None else ""]
                for d in dist]
        _add_table(doc, ["Benefit", "N", "%"], rows)
    if visuals:
        _add_image_or_placeholder(doc, visuals[0])

    s11_2 = sections.get("s11_2", {})
    _add_heading(doc, s11_2.get("label", "Value of Additional Benefits"), level=2)
    _add_paragraph(doc, texts.get("s11_2", ""))

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))


# ---------------------------------------------------------------------------
# Part 12 — Crop Module (renders whenever analysis_results.json has a
# parts.part_12 block -- see run_analysis.py's build_sections(), gated to
# report_scope=="africa"). Every figure here is Vietnam crop clients only.
# ---------------------------------------------------------------------------

def build_part_12(doc, package: dict, texts: dict):
    _add_heading(doc, f"Part 12: {package['title']}", level=1)
    sections = package.get("sections", {})
    visuals  = package.get("visuals", [])

    s12_1 = sections.get("s12_1", {})
    _add_heading(doc, s12_1.get("label", "Weather-Shock Recovery"), level=2)
    _add_paragraph(doc, texts.get("s12_1", ""))
    if visuals:
        _add_image_or_placeholder(doc, visuals[0])

    s12_2 = sections.get("s12_2", {})
    _add_heading(doc, s12_2.get("label", "Farming Approach Change"), level=2)
    _add_paragraph(doc, texts.get("s12_2", ""))

    _add_insight_box(doc, texts.get("insight", ""), sections.get("insight", {}).get("verbatims", []))


# ---------------------------------------------------------------------------
# Main assembly
# ---------------------------------------------------------------------------

def assemble(packages: list, written_texts: dict, run_id: str, output_path: Path):
    doc = Document()
    _set_default_font(doc, "Calibri", 11)
    _apply_brand_heading_color(doc)

    # Cover -- title reflects this run's actual scope, a 4-way split:
    #   1. Single country (meta["country"] != DEFAULT_COUNTRY): that country's label.
    #   2. LARCO regional rollup -- either the legacy dataset_schema == "larco"
    #      export, or report_scope == "lacro" filtering the unified schema
    #      down to LACRO clients (see report_scopes.py and writer.py's
    #      _lacro_scoped()) -- all LACRO countries combined, but NOT the true
    #      global portfolio. Both mechanisms render identically.
    #   3. Any other named report_scope (e.g. "africa" -- see
    #      report_scopes.py): that scope's own label.
    #   4. True Africa+Vietnam global portfolio (dataset_schema ==
    #      "africa_vietnam", no country filter, no report_scope) -- the
    #      original, still-default behavior.
    # Previously only case 1 and the legacy half of case 2 were handled, so a
    # report_scope == "lacro" run (country == "default", dataset_schema ==
    # "africa_vietnam") fell all the way through to "Global Portfolio" here
    # despite being correctly filtered to LACRO clients underneath -- caught
    # from a real generated report titled "Global Portfolio" throughout.
    meta = _load_analysis_meta(run_id)
    period_label = format_period_label(run_id)
    n_total = meta.get("n_total")
    country = meta.get("country")
    report_scope = meta.get("report_scope")

    doc.add_heading("VisionFund International", level=0)
    if country and country != DEFAULT_COUNTRY:
        country_label = meta.get("country_label") or country.title()
        doc.add_heading(f"Insurance Impact Report: {country_label}, {period_label}", level=1)
        if n_total:
            doc.add_paragraph(f"Covering {n_total:,} client responses from {country_label}.")
    elif _lacro_scoped(meta):
        # LACRO_SHORT_LABEL/REPORT_SCOPES (report_scopes.py) are the single
        # source of truth for this region's name -- previously hardcoded
        # here as "LARCO" (the codebase's internal naming convention, not
        # the correct reader-facing spelling), which is exactly how it
        # leaked into the writing LLM's own prose throughout the report
        # (the writer's prompt opens with "You are writing the
        # {report_title}," so a wrong title propagates everywhere).
        doc.add_heading(f"Insurance Impact Report: {LACRO_SHORT_LABEL} Regional Portfolio, {period_label}", level=1)
        if n_total:
            doc.add_paragraph(
                f"Covering {n_total:,} client responses across VisionFund's "
                f"{REPORT_SCOPES['lacro']['label']} insurance portfolio."
            )
    elif report_scope:
        scope_label = meta.get("report_scope_label") or report_scope
        doc.add_heading(f"Insurance Impact Report: {scope_label} Portfolio, {period_label}", level=1)
        if n_total:
            doc.add_paragraph(f"Covering {n_total:,} client responses across VisionFund's {scope_label} insurance portfolio.")
    else:
        doc.add_heading(f"Insurance Impact Report: Global Portfolio, {period_label}", level=1)
        if n_total:
            doc.add_paragraph(f"Covering {n_total:,} client responses across the VisionFund insurance portfolio.")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    doc.add_page_break()

    full_analysis = _load_full_analysis(run_id)
    if full_analysis.get("parts"):
        qual = _load_qualitative_results(run_id)
        _add_executive_summary(doc, full_analysis, qual)
        doc.add_page_break()

    about_survey = _load_analysis_part(run_id, "about_survey")
    if about_survey:
        _add_about_survey_section(
            doc, about_survey,
            data_notes=full_analysis.get("data_notes"),
            quality_flags=full_analysis.get("data_quality_flags"),
        )
        doc.add_page_break()

    # part_8 (Kling Index) has no builder -- intentionally dashboard-only, never
    # part of the written report. See part_8.py's module docstring.
    builders = {
        "part_1": build_part_1,
        "part_2": build_part_2,
        "part_3": build_part_3,
        "part_4": build_part_4,
        "part_5": build_part_5,
        "part_6": build_part_6,
        "part_7": build_part_7,
        "part_9": build_part_9,
        "part_10": build_part_10,
        "part_11": build_part_11,
        "part_12": build_part_12,
    }

    # Part 10 (Trend Comparison, when present) moves to the front of the
    # findings sections -- right after About This Survey, ahead of Part 1 --
    # instead of sitting last after Parts 1-9. It's the headline wave-over-
    # wave result and belongs near the executive summary/about-survey front
    # matter, not buried at the end. sorted() is stable, so every other part
    # keeps its original relative order; only part_10 moves.
    ordered_packages = sorted(packages, key=lambda p: 0 if p["part"] != "part_10" else -1)

    for package in ordered_packages:
        part_key = package["part"]
        texts    = written_texts.get(part_key, {})
        builder  = builders.get(part_key)
        if builder:
            builder(doc, package, texts)
            if texts.get("_generation_failed"):
                log.warning(f"{part_key}: inserting manual-write-up placeholder (generation failed)")
                _add_generation_failure_notice(doc, texts.get("_error", "unknown error"))
            doc.add_page_break()
        else:
            log.warning(f"No builder for {part_key} — skipping")

    part_2_package = next((p for p in packages if p["part"] == "part_2"), None)
    if part_2_package:
        prot_flags = (
            part_2_package.get("sections", {}).get("s2_3", {})
            .get("qualitative", {}).get("protection_flags", [])
        )
        _add_protection_signals_annex(doc, prot_flags, _compute_severity_counts(prot_flags))

    included_parts = {package["part"] for package in packages}
    if "part_4" in included_parts or "part_5" in included_parts:
        _add_methodology_appendix(doc, meta=meta)

    doc.save(str(output_path))
    log.info(f"Report saved: {output_path}")
