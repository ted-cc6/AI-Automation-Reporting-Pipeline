"""
Stage 6 -- assemble.

Builds:
  outputs/GEDSI_Insurance_Report_<date>.docx        -- the report itself
  outputs/GEDSI_Insurance_Supporting_Workbook_<date>.xlsx -- underlying tables
  outputs/run_manifest_<date>.json                  -- auditability record
"""
from __future__ import annotations

import datetime as dt
import json
from ast import literal_eval
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, Pt, RGBColor

from gedsi_pipeline import config, visuals
from gedsi_pipeline.ingest import _read_raw_rows

SECTION_ORDER = [
    "executive_summary",
    "access_understanding",
    "claims_experience",
    "additional_services",
    "wellbeing_resilience",
    "satisfaction_nps",
    "financial_inclusion",
    "disability_crosscut",
    "product_region_notes",
    "external_evidence",
    "benchmarking",
    "recommendations",
]


def _load_json(path):
    return json.loads(path.read_text(encoding="utf-8"))


def _add_demo_table(doc: Document, demo_df: pd.DataFrame, cut_name: str):
    sub = demo_df[demo_df["cut"] == cut_name]
    if sub.empty:
        return
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 2"
    hdr = table.rows[0].cells
    for i, col in enumerate(["Segment", "Female n", "Male n", "Female %", "Pt diff (F-M)"]):
        hdr[i].text = col
    for _, row in sub.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["value"])
        cells[1].text = str(int(row["Female_n"]))
        cells[2].text = str(int(row["Male_n"]))
        cells[3].text = f"{row['Female_pct_of_all_female']:.1f}%" if pd.notna(row["Female_pct_of_all_female"]) else "-"
        pd_val = row["pt_diff"]
        cells[4].text = f"{pd_val:+.1f}" if pd.notna(pd_val) else "-"
    doc.add_paragraph()


def _add_nps_note(doc: Document) -> None:
    """Fixed, code-authored explanation of NPS methodology -- not LLM-authored,
    same reasoning as the Limitations & Methodology section: this needs to be
    exactly right and identically worded every time, not re-derived per run.
    Placed once, directly under the Executive Summary (the report's first and
    most prominent NPS mention), since that's where the ambiguity this note
    heads off actually shows up for a reader."""
    p = doc.add_paragraph()
    run = p.add_run(
        "Note on NPS: a Net Promoter Score is not on the same 0-10 scale as the survey's own "
        "recommendation-likelihood question. NPS is the percentage of respondents who are Promoters "
        "(scored 9-10) minus the percentage who are Detractors (scored 0-6), so it is a percentage-point "
        "gap, ranging from -100 to +100, not a rating out of 10. An NPS of, for example, 20 means the "
        "promoter share exceeds the detractor share by 20 percentage points."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x63, 0x60)


def _add_understanding_chart(doc: Document, df: pd.DataFrame, work_dir: Path) -> None:
    """Code-generated figure, not LLM-authored, for the same reason as
    _add_nps_note: the numbers must exactly match quant_tables/
    understanding_by_product.csv every time, not be redrawn in words by the
    model."""
    chart_path = visuals.render_understanding_by_product_chart(
        df, work_dir / "visuals" / "understanding_by_product_sex.png"
    )
    doc.add_picture(str(chart_path), width=Inches(5.5))
    caption = doc.add_paragraph()
    run = caption.add_run(
        "Figure: share of clients reporting \"very poor\" understanding of their own coverage, by "
        "product and sex. This chart covers clients who already hold each policy; the survey does not "
        "include clients who considered a product and did not enroll, so it cannot show that poor "
        "understanding causes lower enrollment in a product, or the reverse, only that the two differ "
        "together within a specific product."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x63, 0x60)
    doc.add_paragraph()


def _add_section(doc: Document, title: str, headline: str, paragraphs: list[str]):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    run = p.add_run(headline)
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for para_text in paragraphs:
        doc.add_paragraph(para_text)
    doc.add_paragraph()


def build_docx(df: pd.DataFrame, drafts_dir, packs_dir, quant_dir, work_dir: Path, run_stats: dict) -> Document:
    doc = Document()

    ORANGE = RGBColor(0xED, 0x7D, 0x31)  # Word's default "Accent 2" orange
    for style_name in ("Title", "Heading 1", "Heading 2"):
        doc.styles[style_name].font.color.rgb = ORANGE

    title = doc.add_heading("VisionFund 2026 Insurance Survey", level=0)
    sub = doc.add_paragraph()
    sub_run = sub.add_run("GEDSI Analysis Report -- Gender Equality, Disability and Social Inclusion")
    sub_run.bold = True
    sub_run.font.size = Pt(14)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Generated {dt.date.today().isoformat()} | n={run_stats['n_respondents']} respondents"
    ).italic = True
    doc.add_page_break()

    # Executive summary demographic table up front (PDF-specified front-of-report table)
    demo_df = pd.read_csv(quant_dir / "demographic_table.csv")
    exec_draft = _load_json(drafts_dir / "executive_summary.json")
    doc.add_heading("Executive Summary & Demographics", level=1)
    p = doc.add_paragraph()
    run = p.add_run(exec_draft["headline_insight"])
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for para in exec_draft["paragraphs"]:
        doc.add_paragraph(para)
    _add_nps_note(doc)
    doc.add_heading("Survey Demographics by Insurance Type", level=2)
    _add_demo_table(doc, demo_df, "Insurance Type")
    doc.add_heading("Survey Demographics by Country", level=2)
    _add_demo_table(doc, demo_df, "Country")
    doc.add_paragraph()

    section_titles = {
        "access_understanding": "Access & Understanding of Insurance",
        "claims_experience": "Claims Experience",
        "additional_services": "Additional Services & Value-Added Benefits",
        "wellbeing_resilience": "Wellbeing & Financial Resilience",
        "satisfaction_nps": "Client Satisfaction & Net Promoter Score by Gender",
        "financial_inclusion": "Financial Inclusion & First-Time Access",
        "disability_crosscut": "Disability & Vulnerability Cross-Cut",
        "product_region_notes": "Product & Region Notes",
        "external_evidence": "External Evidence & Context",
        "benchmarking": "Benchmarking",
        "recommendations": "Recommendations / Actions",
    }
    for key in SECTION_ORDER:
        if key == "executive_summary":
            continue
        draft = _load_json(drafts_dir / f"{key}.json")
        _add_section(doc, section_titles[key], draft["headline_insight"], draft["paragraphs"])
        if key == "access_understanding":
            _add_understanding_chart(doc, df, work_dir)

    # Limitations & Methodology -- built directly from run facts, not LLM-authored
    doc.add_heading("Limitations & Methodology", level=1)
    limitations = [
        f"Analysis covers {run_stats['n_respondents']} respondents out of {run_stats['n_raw']} raw "
        f"submissions. Excluded: {run_stats['n_test_removed']} test/QA submission(s), "
        f"{run_stats['n_duplicates_removed']} exact-content duplicate submission(s) (one canonical copy "
        f"kept per duplicate group), {run_stats['n_non_consenting_removed']} respondent(s) who declined "
        f"consent, and {run_stats['n_out_of_scope_removed']} respondent(s) outside the study's country "
        f"scope. This screening is aligned with the Cupboard Week insurance report pipeline so both "
        f"reports describe the same population from the same raw export.",
        f"Quantitative comparisons use a minimum cell size of n={config.MIN_N}; comparisons below this "
        f"threshold are marked not-statistically-robust (n_ok=false) and treated as directional only.",
        f"All gender- and disability-comparison p-values in this report were corrected in a single "
        f"Benjamini-Hochberg false discovery rate pass (alpha={config.FDR_ALPHA}) to control for the "
        f"number of comparisons made; only findings surviving this correction are described as significant.",
        "Qualitative themes for the three Net Promoter Score driver questions were induced and coded by "
        f"Claude ({config.CLAUDE_MODEL}) under human review: a human reviewed and edited each theme "
        "codebook before it was applied to the full dataset. Smaller 'please specify' free-text fields "
        "are presented as raw illustrative quotes, not formally coded themes, due to low response counts.",
        "This survey instrument covers insurance clients only. GEDSI report elements specific to "
        "credit/savings products in VisionFund's broader reporting framework (loan goal tracking, "
        "savings group dynamics, household-head/group-leader sub-segments, Empowered Worldview module "
        "attribution) do not apply to this dataset and are omitted rather than approximated.",
        "No internal or external benchmark file (e.g. 60 Decibels, regional/global portfolio benchmarks) "
        "was supplied with this dataset. The Benchmarking section and per-indicator commentary state this "
        "gap explicitly rather than citing unverified figures.",
        "Respondents outside the study's 8-country scope (Rwanda, Ghana, Zambia, Malawi, Uganda, "
        "Tanzania, Kenya, Vietnam) are excluded entirely rather than reported as a small directional "
        "share; Vietnam (Crop Insurance) remains in scope but is still a small share of the sample, so "
        "its country/region cuts are directional, not statistically robust.",
    ]
    for note in limitations:
        doc.add_paragraph(note, style="List Bullet")

    return doc


def build_workbook(quant_dir, work_dir) -> pd.io.excel._OpenpyxlWriter:
    from openpyxl import Workbook
    wb = Workbook()
    wb.remove(wb.active)

    def add_df_sheet(name, df):
        ws = wb.create_sheet(title=name[:31])
        ws.append(list(df.columns))
        for row in df.itertuples(index=False):
            ws.append(list(row))

    for csv_name in ["demographic_table", "gender_comparisons", "disability_comparisons", "nps_by_group"]:
        add_df_sheet(csv_name, pd.read_csv(quant_dir / f"{csv_name}.csv"))

    for q in ["nps_detractor_reasons", "nps_passive_reasons", "nps_promoter_reasons"]:
        path = work_dir / "theme_tables" / f"{q}.csv"
        if path.exists():
            add_df_sheet(f"theme_{q}", pd.read_csv(path))

    quote_bank_path = work_dir / "quote_bank.json"
    if quote_bank_path.exists():
        qb = _load_json(quote_bank_path)
        rows = []
        for question, themes in qb.items():
            for theme, quotes in themes.items():
                for q in quotes:
                    rows.append({"question": question, "theme": theme, **q})
        add_df_sheet("quote_bank", pd.DataFrame(rows))

    return wb


def main(csv_path=None, work_dir: Path | None = None, output_dir: Path | None = None):
    work_dir = work_dir or config.WORK_DIR
    output_dir = output_dir or config.OUTPUT_DIR
    quant_dir = work_dir / "quant_tables"
    packs_dir = work_dir / "evidence_packs"
    drafts_dir = work_dir / "drafts"
    output_dir.mkdir(parents=True, exist_ok=True)

    df = pd.read_parquet(work_dir / "response_frame.parquet")
    _header, raw_rows = _read_raw_rows(csv_path or config.RAW_CSV_PATH)  # csv.reader-based: handles embedded newlines in quoted fields
    n_raw = len(raw_rows)

    screening_summary_path = work_dir / "screening_summary.json"
    if screening_summary_path.exists():
        screen_stats = json.loads(screening_summary_path.read_text(encoding="utf-8"))
    else:
        # Stale work_dir from before per-reason screening existed -- only
        # the raw/final counts are known, not the breakdown by reason.
        screen_stats = {
            "n_test_removed": 0, "n_duplicates_removed": 0,
            "n_non_consenting_removed": n_raw - len(df), "n_out_of_scope_removed": 0,
        }

    run_stats = {
        **screen_stats,
        "n_respondents": len(df),
        "n_raw": n_raw,
    }

    date_str = dt.date.today().isoformat()

    doc = build_docx(df, drafts_dir, packs_dir, quant_dir, work_dir, run_stats)
    docx_path = output_dir / f"GEDSI_Insurance_Report_{date_str}.docx"
    doc.save(docx_path)
    print(f"Wrote {docx_path}")

    wb = build_workbook(quant_dir, work_dir)
    xlsx_path = output_dir / f"GEDSI_Insurance_Supporting_Workbook_{date_str}.xlsx"
    wb.save(xlsx_path)
    print(f"Wrote {xlsx_path}")

    cache_files = list(config.CACHE_DIR.glob("*.json"))
    manifest = {
        "generated_at": dt.datetime.now().isoformat(),
        "model": config.CLAUDE_MODEL,
        "min_n": config.MIN_N,
        "fdr_alpha": config.FDR_ALPHA,
        **run_stats,
        "n_cached_llm_responses": len(cache_files),
        "sections": SECTION_ORDER,
        "outputs": {"docx": str(docx_path), "xlsx": str(xlsx_path)},
    }
    manifest_path = output_dir / f"run_manifest_{date_str}.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {manifest_path}")

    return {"docx_path": docx_path, "xlsx_path": xlsx_path, "manifest_path": manifest_path}


if __name__ == "__main__":
    main()
